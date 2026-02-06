from __future__ import annotations

import os
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Tuple, Optional

import requests
from docx import Document

from config.secure_store import getsecret

PPLX_URL = "https://api.perplexity.ai/chat/completions"
DEFAULT_MODEL = "sonar"
DEFAULT_TIMEOUT_S = 60
MAX_RETRIES = 3


@dataclass
class SummaryResult:
    ok: bool
    text: str = ""
    error: str = ""
    status_code: Optional[int] = None


def _write_text(path: Path, content: str) -> None:
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
    except Exception:
        pass


def _is_pplx_key(v: str) -> bool:
    v = (v or "").strip()
    return v.startswith("pplx-") and len(v) >= 20


def _get_perplexity_key(cfg: Dict[str, Any]) -> Tuple[str, str]:
    tried = []

    secure_keys = [
        "perplexity_api_key",
        "perplexity_key",
        "pplx_api_key",
        "perplexity_token",
        "pplx_token",
        "perplexityapikey",
        "perplexity",
        "pplx",
        "api_perplexity_key",
    ]
    for k in secure_keys:
        tried.append(f"secure_store:{k}")
        v = (getsecret(cfg, k) or "").strip()
        if _is_pplx_key(v):
            return v, f"secure_store:{k}"

    cfg_keys = secure_keys
    for k in cfg_keys:
        tried.append(f"cfg:{k}")
        v = str((cfg or {}).get(k, "")).strip()
        if _is_pplx_key(v):
            return v, f"cfg:{k}"

    env_keys = ["PERPLEXITY_API_KEY", "PPLX_API_KEY", "PERPLEXITY_KEY", "PPLX_KEY"]
    for k in env_keys:
        tried.append(f"env:{k}")
        v = (os.getenv(k) or "").strip()
        if _is_pplx_key(v):
            return v, f"env:{k}"

    for ck, cv in (cfg or {}).items():
        try:
            s = str(cv).strip()
        except Exception:
            continue
        if _is_pplx_key(s):
            tried.append(f"cfg:*looks_like_key:{ck}")
            return s, f"cfg:*looks_like_key:{ck}"

    return "", f"not_found; tried={', '.join(tried)}"


def _call_perplexity(api_key: str, prompt: str, model: str, timeout_s: int = DEFAULT_TIMEOUT_S) -> SummaryResult:
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
    }

    last_err = ""
    last_code = None

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            r = requests.post(PPLX_URL, headers=headers, json=payload, timeout=timeout_s)
            last_code = r.status_code
            if r.status_code == 200:
                data = r.json()
                text = ((((data.get("choices") or [{}])[0]).get("message") or {}).get("content") or "").strip()
                if text:
                    return SummaryResult(ok=True, text=text, status_code=r.status_code)
                return SummaryResult(ok=False, error="Réponse Perplexity vide.", status_code=r.status_code)

            if r.status_code in (429, 500, 502, 503, 504):
                last_err = f"HTTP {r.status_code}: {r.text[:800]}"
                time.sleep(min(2 ** attempt, 8))
                continue

            return SummaryResult(ok=False, error=f"HTTP {r.status_code}: {r.text[:2000]}", status_code=r.status_code)

        except Exception as e:
            last_err = repr(e)
            time.sleep(min(2 ** attempt, 8))

    return SummaryResult(ok=False, error=last_err or "Erreur inconnue Perplexity.", status_code=last_code)


def generate_meeting_docx(transcript_path: Path, session_dir: Path, cfg: Dict[str, Any]) -> Path:
    transcript_path = Path(transcript_path)
    session_dir = Path(session_dir)
    session_dir.mkdir(parents=True, exist_ok=True)

    try:
        transcript_text = transcript_path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        transcript_text = f"(Impossible de lire la transcription: {e!r})"

    key, key_src = _get_perplexity_key(cfg or {})

    pplx_err = ""
    summary_text = ""

    if not bool((cfg or {}).get("postprocess_enable_perplexity_summary", True)):
        pplx_err = "Résumé Perplexity désactivé dans les réglages."
    elif not key:
        pplx_err = f"clé Perplexity absente (source={key_src})"
    else:
        model = str((cfg or {}).get("perplexity_model") or DEFAULT_MODEL).strip()
        template = str((cfg or {}).get("summary_template") or "Compte rendu pro")
        template_hint = ""
        if template == "Vidéo YouTube":
            template_hint = "Le compte rendu doit être adapté à une vidéo YouTube (ton pédagogique, points clés, conclusions)."
        elif template == "Webinaire":
            template_hint = "Le compte rendu doit être adapté à un webinaire (structure pédagogique, questions/réponses, points clés)."
        elif template == "Réunion Discord (asso)":
            template_hint = "Le compte rendu doit être adapté à une réunion Discord associative (actions concrètes, décisions, suivi)."
        elif template == "Recrutement":
            template_hint = "Le compte rendu doit être adapté à un entretien de recrutement (profil, compétences, points forts/faiblesses)."
        else:
            template_hint = "Le compte rendu doit être professionnel et synthétique."
        prompt = (
            "Tu es un assistant chargé de produire un compte rendu fidèle et exploitable.\n"
            "Le contexte peut être professionnel ou personnel.\n\n"
            f"{template_hint}\n\n"
            "À partir de la transcription ci-dessous, produis STRICTEMENT la structure suivante :\n\n"
            "=== RÉSUMÉ GLOBAL ===\n"
            "Synthèse claire de la situation en 5 à 10 lignes.\n\n"
            "=== POINTS ÉVOQUÉS ===\n"
            "- Liste factuelle des sujets abordés.\n\n"
            "=== DÉCISIONS ===\n"
            "- Décisions prises explicitement (sinon écrire : Aucune décision formalisée).\n\n"
            "=== ACTIONS À FAIRE ===\n"
            "- Action | Responsable | Échéance | Priorité\n"
            "- Si une information est absente, écrire 'Non précisée'.\n\n"
            "=== DEADLINES / ÉCHÉANCES ===\n"
            "- Récapitulatif des dates mentionnées (ou 'Aucune').\n\n"
            "=== RISQUES / POINTS EN SUSPENS ===\n"
            "- Points non tranchés, ambiguïtés, sujets à clarifier.\n\n"
            "Règles importantes :\n"
            "- Ne pas inventer d'informations.\n"
            "- Être concis, structuré, professionnel.\n"
            "- Utiliser un français clair.\n\n"
            "TRANSCRIPTION :\n"
            f"""{transcript_text}"""
        )
        res = _call_perplexity(key.strip(), prompt, model=model)
        if res.ok:
            summary_text = res.text.strip()
        else:
            pplx_err = res.error or "Erreur inconnue Perplexity."

    if pplx_err:
        _write_text(session_dir / "perplexity_error.txt", pplx_err)

    doc = Document()
    doc.add_heading("Résumé de Réunion", level=1)
    if summary_text:
        doc.add_paragraph(summary_text)
    else:
        doc.add_paragraph("Résumé non disponible (clé Perplexity absente ou erreur d'appel).")

    if bool((cfg or {}).get("postprocess_extract_participants", False)):
        try:
            names = extract_participant_names(transcript_path, cfg)
            if names:
                doc.add_heading("Participants (IA)", level=2)
                for line in names.splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip())
        except Exception:
            pass

    doc.add_heading("Transcription", level=1)
    for line in transcript_text.splitlines():
        if line.strip():
            doc.add_paragraph(line)

    out_path = session_dir / "Résumé de Réunion.docx"
    doc.save(str(out_path))
    return out_path


def generate_subject_from_transcript(transcript_path: Path, cfg: Dict[str, Any]) -> str:
    try:
        transcript_text = Path(transcript_path).read_text(encoding="utf-8", errors="replace")
    except Exception:
        transcript_text = ""
    key, _ = _get_perplexity_key(cfg or {})
    if not key or not transcript_text.strip():
        return ""
    prompt = (
        "Tu dois proposer un sujet très court pour une réunion.\n"
        "Donne uniquement un titre court (4 à 8 mots), sans guillemets.\n\n"
        "TRANSCRIPTION:\n"
        f"{transcript_text}"
    )
    res = _call_perplexity(key.strip(), prompt, model=str((cfg or {}).get("perplexity_model") or DEFAULT_MODEL).strip())
    if res.ok:
        return (res.text or "").strip().strip('"')
    return ""


def extract_participant_names(transcript_path: Path, cfg: Dict[str, Any]) -> str:
    try:
        transcript_text = Path(transcript_path).read_text(encoding="utf-8", errors="replace")
    except Exception:
        transcript_text = ""
    key, _ = _get_perplexity_key(cfg or {})
    if not key or not transcript_text.strip():
        return ""
    prompt = (
        "Analyse la transcription et tente d'identifier les noms des participants.\n"
        "Retourne une liste simple, une ligne par participant (Nom - rôle si possible).\n"
        "Si aucune info fiable, réponds 'Non précisé'.\n\n"
        "TRANSCRIPTION:\n"
        f"{transcript_text}"
    )
    res = _call_perplexity(key.strip(), prompt, model=str((cfg or {}).get("perplexity_model") or DEFAULT_MODEL).strip())
    if res.ok:
        return (res.text or "").strip()
    return ""

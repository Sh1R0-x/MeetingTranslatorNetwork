from __future__ import annotations

import json
import os
import queue
import re
import shutil
import subprocess
import traceback
import wave
import html
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional
import time

try:
    import sounddevice as sd
except Exception:  # pragma: no cover - optional at runtime
    sd = None

from PyQt6.QtCore import QEvent, QTimer, Qt, pyqtSignal, QSize, QThread, QPropertyAnimation, QEasingCurve
from PyQt6.QtGui import QGuiApplication, QTextCursor, QTextCharFormat, QFont, QColor
from PyQt6.QtGui import QAction
from PyQt6.QtWidgets import QMenu
from PyQt6.QtWidgets import (
    QAbstractItemView,
    QApplication,
    QCheckBox,
    QComboBox,
    QDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QInputDialog,
    QPlainTextEdit,
    QTextEdit,
    QTextBrowser,
    QProgressBar,
    QPushButton,
    QGraphicsOpacityEffect,
    QSizePolicy,
    QSplitter,
    QTabWidget,
    QTableWidget,
    QTableWidgetItem,
    QToolButton,
    QVBoxLayout,
    QWidget,
    QFrame,
)

import common as common_mod
from config.secure_store import load_config, save_config, getsecret
from common import LOG_PATH
from services.recorder_service import RecorderService

APP_NAME = "MeetingTranslatorNetwork"
APP_VERSION = "2026"
DEFAULT_SESSIONS_DIR = str((Path.home() / "Documents" / APP_NAME / "recordings"))
OPENAI_TRANSCRIBE_PER_MIN = 0.006
ASSEMBLYAI_STREAMING_PER_HOUR = 0.15
ASSEMBLYAI_STREAMING_PER_MIN = ASSEMBLYAI_STREAMING_PER_HOUR / 60.0
DEEPGRAM_STREAMING_NOVA3_PER_MIN = 0.0077


def log_line(msg: str):
    try:
        common_mod.log_line(msg)
    except Exception:
        pass


@dataclass
class LiveMessage:
    timestamp: str
    source: str
    source_raw: str = ""
    ts_sec: float = 0.0
    text_en: str = ""
    text_fr: str = ""
    important: bool = False


class StatusPill(QToolButton):
    def __init__(self, text: str, parent=None):
        super().__init__(parent)
        self._base_text = str(text)
        self.setText(f"{self._base_text} OFF")
        self.setCheckable(True)
        self.setChecked(False)
        self.setProperty("kind", "segment")
        self.setProperty("state", "off")
        self.setProperty("selected", "false")
        self.setFocusPolicy(Qt.FocusPolicy.NoFocus)

    def set_state(self, state: str):
        state = str(state or "off")
        self.setProperty("state", state)
        is_on = state in ("ok", "warn")
        self.setText(f"{self._base_text} {'ON' if is_on else 'OFF'}")
        self.setProperty("selected", "true" if is_on else "false")
        self.setChecked(is_on)
        self.style().polish(self)


class PulseButton(QPushButton):
    """
    Small UI affordance: visible feedback even when the underlying action has no immediate UI change.
    """

    def __init__(self, text: str = "", parent=None):
        super().__init__(text, parent)
        self._pulse_fx = QGraphicsOpacityEffect(self)
        self._pulse_fx.setOpacity(1.0)
        self.setGraphicsEffect(self._pulse_fx)
        self._pulse_anim = QPropertyAnimation(self._pulse_fx, b"opacity", self)
        self._pulse_anim.setDuration(420)
        self._pulse_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.clicked.connect(self._pulse)

    def _pulse(self):
        # Restart a short pulse animation on click.
        try:
            self._pulse_anim.stop()
            self._pulse_anim.setStartValue(1.0)
            self._pulse_anim.setKeyValueAt(0.5, 0.65)
            self._pulse_anim.setEndValue(1.0)
            self._pulse_anim.start()
        except Exception:
            pass

        # Short "flash" state for clearer click feedback (background tint), then revert.
        try:
            self.setProperty("flash", "true")
            self.style().polish(self)

            def _clear():
                try:
                    self.setProperty("flash", "false")
                    self.style().polish(self)
                except Exception:
                    pass

            # Slightly longer flash for a smoother, more visible click feedback.
            QTimer.singleShot(720, _clear)
        except Exception:
            pass

        # Make sure the flash is painted even if a heavy handler runs right after the click.
        try:
            QApplication.processEvents()
        except Exception:
            pass


class LiveMessageWidget(QWidget):
    def __init__(self, msg: LiveMessage, parent=None):
        super().__init__(parent)
        self.msg = msg
        self.setObjectName("LiveMessageCard")

        self.root = QVBoxLayout(self)
        self.root.setContentsMargins(16, 12, 16, 12)
        self.root.setSpacing(8)

        meta = QHBoxLayout()
        meta.setSpacing(8)

        self.lbl_time = QLabel(msg.timestamp)
        self.lbl_time.setObjectName("LiveMetaTime")
        self.lbl_time.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
        meta.addWidget(self.lbl_time)

        self.lbl_source = QLabel(msg.source)
        self.lbl_source.setObjectName("LiveMetaSource")
        self.lbl_source.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
        meta.addWidget(self.lbl_source)
        meta.addStretch(1)

        self.root.addLayout(meta)

        self.en_bubble = QWidget()
        self.en_bubble.setObjectName("BubbleEN")
        self.en_bubble.setSizePolicy(QSizePolicy.Policy.Maximum, QSizePolicy.Policy.Preferred)
        en_layout = QHBoxLayout(self.en_bubble)
        en_layout.setContentsMargins(12, 8, 12, 8)
        en_layout.setSpacing(8)
        self.tag_en = QLabel("Anglais")
        self.tag_en.setObjectName("LangTagEN")
        self.tag_en.setVisible(False)
        self.txt_en = QLabel(msg.text_en)
        self.txt_en.setWordWrap(True)
        self.txt_en.setObjectName("LiveTextEN")
        self.txt_en.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.txt_en.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Preferred)
        self.txt_en.setMinimumWidth(0)
        en_layout.addWidget(self.tag_en)
        en_layout.addWidget(self.txt_en)

        self.fr_bubble = QWidget()
        self.fr_bubble.setObjectName("BubbleFR")
        self.fr_bubble.setSizePolicy(QSizePolicy.Policy.Maximum, QSizePolicy.Policy.Preferred)
        fr_layout = QHBoxLayout(self.fr_bubble)
        fr_layout.setContentsMargins(12, 8, 12, 8)
        fr_layout.setSpacing(8)
        self.tag_fr = QLabel("Francais")
        self.tag_fr.setObjectName("LangTagFR")
        self.tag_fr.setVisible(False)
        self.txt_fr = QLabel(msg.text_fr)
        self.txt_fr.setWordWrap(True)
        self.txt_fr.setObjectName("LiveTextFR")
        self.txt_fr.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.txt_fr.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Preferred)
        self.txt_fr.setMinimumWidth(0)
        fr_layout.addWidget(self.tag_fr)
        fr_layout.addWidget(self.txt_fr)

        self.en_wrap = QHBoxLayout()
        self.en_wrap.setContentsMargins(0, 0, 0, 0)
        self.en_wrap.addWidget(self.en_bubble, 0, Qt.AlignmentFlag.AlignLeft)
        self.root.addLayout(self.en_wrap)

        self.fr_wrap = QHBoxLayout()
        self.fr_wrap.setContentsMargins(0, 0, 0, 0)
        self.fr_wrap.addWidget(self.fr_bubble, 0, Qt.AlignmentFlag.AlignLeft)
        self.root.addLayout(self.fr_wrap)

        self._content_width = None
        self.update_message(msg)

    def set_speaker_break(self, enabled: bool):
        if enabled:
            self.root.setContentsMargins(16, 20, 16, 12)
        else:
            self.root.setContentsMargins(16, 12, 16, 12)

    def update_message(self, msg: LiveMessage):
        self.msg = msg
        self.lbl_time.setText(msg.timestamp)
        self.lbl_source.setText(msg.source)
        self.lbl_source.setProperty("source", (msg.source or "").lower())
        self.lbl_source.style().polish(self.lbl_source)
        en = (msg.text_en or "").strip()
        fr = (msg.text_fr or "").strip()
        self.txt_en.setText(en)
        self.txt_fr.setText(fr)
        self.en_bubble.setVisible(bool(en))
        self.fr_bubble.setVisible(bool(fr))
        self.setProperty("important", bool(msg.important))
        self.style().polish(self)
        if self._content_width:
            self._apply_content_width(self._content_width)

    def set_content_width(self, width: int):
        self._content_width = int(width)
        self._apply_content_width(self._content_width)

    def _apply_content_width(self, width: int):
        # Keep bubbles readable and force proper wrapping.
        inner = max(360, int(width) - 24)
        for bubble in (self.en_bubble, self.fr_bubble):
            bubble.setFixedWidth(inner)
        for lbl in (self.txt_en, self.txt_fr):
            lbl.setFixedWidth(max(120, inner - 24))
            lbl.setMinimumWidth(0)
            lbl.adjustSize()


class LiveChatList(QListWidget):
    scrolled = pyqtSignal()
    resized = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("LiveChatList")
        self.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.setVerticalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.viewport().installEventFilter(self)

    def eventFilter(self, obj, event):
        if event.type() == QEvent.Type.Wheel:
            self.scrolled.emit()
        return super().eventFilter(obj, event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.resized.emit()


class LiveChatView(QTextBrowser):
    """
    QTextBrowser-based "chat" view.

    Why: QListWidget + setItemWidget is fragile for dynamic word-wrapped items on Windows;
    QTextBrowser is far more stable for sizing/wrapping and keeps text selectable/copiable.
    """

    scrolled = pyqtSignal()
    resized = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("LiveChatView")
        self.setReadOnly(True)
        self.setOpenExternalLinks(False)
        self.setUndoRedoEnabled(False)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        # Keep a stable viewport width (prevents left/right text shifts when scrollbar appears).
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        self.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        self.setStyleSheet(
            "QTextBrowser#LiveChatView{"
            "background:#0e0e12;"
            "border:1px solid #2a2a38;"
            "border-radius:12px;"
            "}"
        )

        self._messages: list[LiveMessage] = []
        self._in_width_update = False

        # Coalesce frequent updates (streaming partials) into a single re-render.
        self._render_timer = QTimer(self)
        self._render_timer.setSingleShot(True)
        self._render_timer.timeout.connect(self._render_now)

        self._pending_keep_scroll = False
        self._pending_scroll_bottom = False
        self._pending_scroll_value = 0
        self._session_marker_inserted = False

        # Detect user scroll even when they drag the scrollbar.
        bar = self.verticalScrollBar()
        try:
            bar.sliderMoved.connect(lambda _v: self.scrolled.emit())
            bar.actionTriggered.connect(lambda _a: self.scrolled.emit())
        except Exception:
            pass

        # Text formats used by minimal fallback rendering.
        self._fmt_meta = QTextCharFormat()
        self._fmt_meta.setForeground(QColor("#71717a"))
        meta_font = QFont("Inter")
        meta_font.setPointSizeF(9.0)
        meta_font.setWeight(QFont.Weight.Medium)
        self._fmt_meta.setFont(meta_font)

        self._fmt_body = QTextCharFormat()
        self._fmt_body.setForeground(QColor("#e2e8f0"))
        body_font = QFont("Inter")
        body_font.setPointSizeF(11.0)
        self._fmt_body.setFont(body_font)

        # Initialize an empty document.
        self.clear()
        self._update_reading_width()

    def wheelEvent(self, event):
        self.scrolled.emit()
        return super().wheelEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        # Avoid recursive resize storms (Qt may trigger relayouts while we adjust margins/text width).
        QTimer.singleShot(0, self._update_reading_width)
        self.resized.emit()

    def clear_messages(self):
        self._messages.clear()
        self._session_marker_inserted = False
        self.clear()
        self._update_reading_width()

    def request_render(self, *, keep_scroll: bool, scroll_to_bottom: bool):
        self._schedule_render(keep_scroll=keep_scroll, scroll_to_bottom=scroll_to_bottom)

    def add_message(self, msg: LiveMessage, *, keep_scroll: bool, scroll_to_bottom: bool):
        self._messages.append(msg)

        # Incremental append is much faster than re-rendering the whole document.
        if len(self._messages) == 1 and not self._session_marker_inserted:
            self._append_session_marker(msg)
        self._append_message_to_doc(msg)

        if scroll_to_bottom:
            self.scrollToBottom()

    def update_last_message(self, msg: LiveMessage, *, keep_scroll: bool, scroll_to_bottom: bool):
        if self._messages:
            self._messages[-1] = msg
        else:
            self._messages.append(msg)
        # Updating already-rendered rich text reliably is hard; re-render (throttled).
        self._schedule_render(keep_scroll=keep_scroll, scroll_to_bottom=scroll_to_bottom)

    def get_last_message(self) -> Optional[LiveMessage]:
        return self._messages[-1] if self._messages else None

    def scrollToBottom(self):
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.setTextCursor(cursor)
        self.ensureCursorVisible()

    def _schedule_render(self, *, keep_scroll: bool, scroll_to_bottom: bool):
        self._pending_keep_scroll = bool(keep_scroll)
        self._pending_scroll_bottom = bool(scroll_to_bottom)
        try:
            self._pending_scroll_value = int(self.verticalScrollBar().value())
        except Exception:
            self._pending_scroll_value = 0
        if not self._render_timer.isActive():
            # 80ms keeps UI responsive while avoiding a reflow on every partial token.
            self._render_timer.start(80)

    def _render_now(self):
        bar = self.verticalScrollBar()
        saved = self._pending_scroll_value

        self._render_all_messages()

        def _after():
            try:
                if self._pending_scroll_bottom:
                    self.scrollToBottom()
                    return
                if self._pending_keep_scroll:
                    bar.setValue(min(int(saved), bar.maximum()))
            except Exception:
                pass

        QTimer.singleShot(0, _after)

    def _append_message_to_doc(self, msg: LiveMessage):
        try:
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            self._insert_one(cursor, msg)
            self.setTextCursor(cursor)
        except Exception:
            # Fallback: full re-render on unexpected text failures.
            self._schedule_render(keep_scroll=True, scroll_to_bottom=False)

    def _render_all_messages(self):
        try:
            self.clear()
            self._update_reading_width()
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            if self._messages:
                self._insert_session_marker(cursor, self._messages[0])
                self._session_marker_inserted = True
            # Limit to keep UI responsive on very long sessions.
            max_render = 2500
            for msg in (self._messages[-max_render:] if len(self._messages) > max_render else self._messages):
                self._insert_one(cursor, msg)
            self.setTextCursor(cursor)
        except Exception:
            pass

    def _insert_one(self, cursor: QTextCursor, msg: LiveMessage):
        ts = html.escape((msg.timestamp or "").strip())
        src = self._normalize_source_label(msg.source)
        src_html = html.escape(src)

        parts: list[str] = []
        if (msg.text_en or "").strip():
            parts.append((msg.text_en or "").strip())
        if (msg.text_fr or "").strip():
            text_fr = (msg.text_fr or "").strip()
            if not parts or text_fr != parts[-1]:
                parts.append(text_fr)
        body = "\n".join(parts).strip()
        if not body:
            body = "..."

        body_html = html.escape(body).replace("\n", "<br/>")
        # Speaker color palette
        speaker_colors = {
            "MOI": "#0d9af2",
            "PARTICIPANTS": "#8b5cf6",
        }
        source_color = speaker_colors.get(src.upper(), "#8b5cf6")
        block_html = (
            "<div style='margin:0 0 18px 0;'>"
            f"<div style='margin:0 0 6px 0;'>"
            f"<span style='font-size:10px;color:#52525b;font-family:JetBrains Mono,Consolas,monospace;'>[{ts}]</span> "
            f"<span style='font-size:11px;color:{source_color};font-weight:700;font-family:Inter,sans-serif;'>{src_html}</span>"
            f"</div>"
            "<div style='margin-left:60px;background:#13131a;border:1px solid #2a2a38;border-radius:10px;padding:10px 14px;max-width:920px;'>"
            f"<span style='font-size:13px;line-height:1.6;color:#e2e8f0;font-family:Inter,sans-serif;'>{body_html}</span>"
            "</div>"
            "</div>"
        )
        cursor.insertHtml(block_html)
        cursor.insertBlock()

    def _append_session_marker(self, msg: LiveMessage):
        try:
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            self._insert_session_marker(cursor, msg)
            self.setTextCursor(cursor)
            self._session_marker_inserted = True
        except Exception:
            pass

    def _insert_session_marker(self, cursor: QTextCursor, msg: LiveMessage):
        ts = html.escape((msg.timestamp or "").strip() or "--:--")
        marker_html = (
            "<div style='text-align:center;margin:10px 0 20px 0;'>"
            "<span style='display:inline-block;background:#13131a;border:1px solid #2a2a38;"
            "border-radius:999px;padding:6px 16px;font-size:11px;color:#71717a;"
            "font-family:Inter,sans-serif;font-weight:500;letter-spacing:0.3px;'>"
            f"Séance démarrée à {ts}"
            "</span></div>"
        )
        cursor.insertHtml(marker_html)
        cursor.insertBlock()

    def _normalize_source_label(self, source: Optional[str]) -> str:
        raw = (source or "").strip()
        low = raw.lower()
        if low == "participants":
            return "PARTICIPANTS"
        if low == "moi":
            return "MOI"
        if raw:
            return raw
        return "SPEAKER"

    def _update_reading_width(self):
        # Keep a readable measure (avoid full-width lines on very wide windows).
        if self._in_width_update:
            return
        self._in_width_update = True
        try:
            viewport_w = max(1, int(self.viewport().width()))
            max_text_w = 920
            text_w = min(max_text_w, max(520, viewport_w - 28))
            # Keep fixed side margins to avoid visual left/right drift while streaming.
            side = 14
            self.setViewportMargins(side, 10, side, 10)
            self.document().setDocumentMargin(0.0)
            self.document().setTextWidth(float(text_w))
        except Exception:
            pass
        finally:
            self._in_width_update = False


class PopoutLiveChatView(QTextBrowser):
    """Dedicated popout chat view: black background + white bubbles."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("PopoutLiveChatView")
        self.setReadOnly(True)
        self.setOpenExternalLinks(False)
        self.setUndoRedoEnabled(False)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        self.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.setStyleSheet(
            "QTextBrowser#PopoutLiveChatView{"
            "background:#0e0e12;border:1px solid #2a2a38;border-radius:12px;"
            "}"
        )

        self._messages: list[LiveMessage] = []
        self._render_timer = QTimer(self)
        self._render_timer.setSingleShot(True)
        self._render_timer.timeout.connect(self._render_all)
        self._pending_keep_scroll = False
        self._pending_scroll_bottom = False
        self._pending_scroll_value = 0
        self._fmt_meta = QTextCharFormat()
        self._fmt_meta.setForeground(QColor("#71717a"))
        meta_font = QFont("Inter")
        meta_font.setPointSizeF(9.0)
        meta_font.setWeight(QFont.Weight.Medium)
        self._fmt_meta.setFont(meta_font)
        self._fmt_body = QTextCharFormat()
        self._fmt_body.setForeground(QColor("#e2e8f0"))
        body_font = QFont("Inter")
        body_font.setPointSizeF(12.0)
        self._fmt_body.setFont(body_font)

    def clear_messages(self):
        self._messages.clear()
        self.clear()

    def add_message(self, msg: LiveMessage, *, keep_scroll: bool, scroll_to_bottom: bool):
        self._messages.append(msg)
        self._append_message_to_doc(msg)
        if scroll_to_bottom:
            self.scroll_to_bottom()
        elif keep_scroll:
            bar = self.verticalScrollBar()
            bar.setValue(min(bar.value(), bar.maximum()))

    def update_last_message(self, msg: LiveMessage, *, keep_scroll: bool, scroll_to_bottom: bool):
        if self._messages:
            self._messages[-1] = msg
        else:
            self._messages.append(msg)
        self._schedule_render(keep_scroll=keep_scroll, scroll_to_bottom=scroll_to_bottom)

    def set_messages(self, messages: list[LiveMessage]):
        self._messages = list(messages or [])
        self._render_all()

    def _schedule_render(self, *, keep_scroll: bool, scroll_to_bottom: bool):
        self._pending_keep_scroll = bool(keep_scroll)
        self._pending_scroll_bottom = bool(scroll_to_bottom)
        try:
            self._pending_scroll_value = int(self.verticalScrollBar().value())
        except Exception:
            self._pending_scroll_value = 0
        if not self._render_timer.isActive():
            self._render_timer.start(80)

    def _render_all(self):
        bar = self.verticalScrollBar()
        saved = self._pending_scroll_value
        try:
            self.clear()
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            for msg in self._messages[-2500:]:
                self._insert_one(cursor, msg)
            self.setTextCursor(cursor)
        except Exception:
            return

        def _after():
            try:
                if self._pending_scroll_bottom:
                    self.scroll_to_bottom()
                elif self._pending_keep_scroll:
                    bar.setValue(min(int(saved), bar.maximum()))
            except Exception:
                pass

        QTimer.singleShot(0, _after)

    def _append_message_to_doc(self, msg: LiveMessage):
        try:
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)
            self._insert_one(cursor, msg)
            self.setTextCursor(cursor)
        except Exception:
            self._schedule_render(keep_scroll=True, scroll_to_bottom=False)

    def _insert_one(self, cursor: QTextCursor, msg: LiveMessage):
        ts = html.escape((msg.timestamp or "").strip())
        src = self._source_label(msg.source)
        src_html = html.escape(src)
        text_parts = []
        if (msg.text_en or "").strip():
            text_parts.append((msg.text_en or "").strip())
        if (msg.text_fr or "").strip():
            fr = (msg.text_fr or "").strip()
            if not text_parts or fr != text_parts[-1]:
                text_parts.append(fr)
        body = "\n".join(text_parts).strip()
        if not body:
            body = "..."
        body_html = html.escape(body).replace("\n", "<br/>")
        speaker_colors = {
            "MOI": "#0d9af2",
            "PARTICIPANTS": "#8b5cf6",
        }
        source_color = speaker_colors.get(src.upper(), "#8b5cf6")
        block_html = (
            "<div style='margin:0 0 18px 0;'>"
            f"<div style='margin:0 0 6px 0;'>"
            f"<span style='font-size:10px;color:#52525b;font-family:JetBrains Mono,Consolas,monospace;'>[{ts}]</span> "
            f"<span style='font-size:11px;color:{source_color};font-weight:700;font-family:Inter,sans-serif;'>{src_html}</span>"
            f"</div>"
            "<div style='margin-left:60px;background:#13131a;border:1px solid #2a2a38;border-radius:10px;padding:10px 14px;max-width:980px;'>"
            f"<span style='font-size:13px;line-height:1.6;color:#e2e8f0;font-family:Inter,sans-serif;'>{body_html}</span>"
            "</div>"
            "</div>"
        )
        cursor.insertHtml(block_html)
        cursor.insertBlock()

    def _source_label(self, source: Optional[str]) -> str:
        raw = (source or "").strip()
        low = raw.lower()
        if low == "participants":
            return "PARTICIPANTS"
        if low == "moi":
            return "MOI"
        return raw if raw else "SPEAKER"

    def scroll_to_bottom(self):
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.setTextCursor(cursor)
        self.ensureCursorVisible()


class LiveChatWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Chat Live")
        self.resize(900, 700)
        root = QWidget()
        layout = QVBoxLayout(root)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(8)
        self.chk_autoscroll = QCheckBox("Auto-scroll")
        self.chk_autoscroll.setChecked(True)
        layout.addWidget(self.chk_autoscroll, alignment=Qt.AlignmentFlag.AlignLeft)
        self.chat = PopoutLiveChatView()
        layout.addWidget(self.chat, 1)
        self.setCentralWidget(root)

    def add_message(self, msg: LiveMessage, *, keep_scroll: bool):
        auto = bool(self.chk_autoscroll.isChecked())
        self.chat.add_message(msg, keep_scroll=keep_scroll, scroll_to_bottom=auto and not keep_scroll)

    def update_last_message(self, msg: LiveMessage, *, keep_scroll: bool):
        auto = bool(self.chk_autoscroll.isChecked())
        self.chat.update_last_message(msg, keep_scroll=keep_scroll, scroll_to_bottom=auto and not keep_scroll)

    def set_messages(self, messages: list[LiveMessage]):
        self.chat.set_messages(messages)
        if self.chk_autoscroll.isChecked():
            self.chat.scroll_to_bottom()

    def clear_messages(self):
        self.chat.clear_messages()


class TopBar(QWidget):
    start_clicked = pyqtSignal()
    pause_clicked = pyqtSignal()
    stop_clicked = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("TopBar")
        self.setProperty("kind", "panel")
        self.setFixedHeight(64)

        root = QHBoxLayout(self)
        root.setContentsMargins(12, 8, 12, 8)
        root.setSpacing(10)

        left = QHBoxLayout()
        left.setSpacing(10)

        self.timer_badge = QFrame()
        self.timer_badge.setObjectName("TimerBadge")
        self.timer_badge.setProperty("kind", "card")
        badge_layout = QHBoxLayout(self.timer_badge)
        badge_layout.setContentsMargins(10, 6, 10, 6)
        badge_layout.setSpacing(8)
        self.lbl_rec = QLabel("● REC")
        self.lbl_rec.setObjectName("RecLabel")
        self.lbl_timer = QLabel("00:00:00")
        self.lbl_timer.setObjectName("TopBarTimer")
        badge_layout.addWidget(self.lbl_rec)
        badge_layout.addWidget(self.lbl_timer)
        left.addWidget(self.timer_badge)

        divider = QFrame()
        divider.setObjectName("TopDivider")
        divider.setFixedSize(1, 30)
        left.addWidget(divider)

        self.btn_start = PulseButton("Démarrer")
        self.btn_start.setObjectName("BtnStart")
        self.btn_start.setProperty("variant", "primary")
        self.btn_start.setProperty("tone", "success")
        self.btn_pause = PulseButton("Pause")
        self.btn_pause.setObjectName("BtnPause")
        self.btn_pause.setProperty("variant", "ghost")
        self.btn_pause.setProperty("tone", "warning")
        self.btn_stop = PulseButton("Arrêter")
        self.btn_stop.setObjectName("BtnStop")
        self.btn_stop.setProperty("variant", "ghost")
        self.btn_stop.setProperty("tone", "danger")
        self.btn_reset = PulseButton("Reset")
        self.btn_reset.setObjectName("BtnReset")
        self.btn_reset.setProperty("variant", "ghost")
        self.btn_reset.setProperty("tone", "neutral")
        self.btn_pause.setEnabled(False)
        self.btn_stop.setEnabled(False)
        left.addWidget(self.btn_start)
        left.addWidget(self.btn_pause)
        left.addWidget(self.btn_stop)
        left.addWidget(self.btn_reset)
        root.addLayout(left)
        root.addStretch(1)

        # Subtle recording animation (pulsing REC) for better affordance.
        self._rec_fx = QGraphicsOpacityEffect(self.lbl_rec)
        self._rec_fx.setOpacity(1.0)
        self.lbl_rec.setGraphicsEffect(self._rec_fx)
        self._rec_anim = QPropertyAnimation(self._rec_fx, b"opacity", self)
        self._rec_anim.setDuration(900)
        self._rec_anim.setEasingCurve(QEasingCurve.Type.InOutSine)
        self._rec_anim.setStartValue(1.0)
        self._rec_anim.setKeyValueAt(0.5, 0.35)
        self._rec_anim.setEndValue(1.0)
        self._rec_anim.setLoopCount(-1)

        right = QHBoxLayout()
        right.setSpacing(10)
        self.segment_group = QFrame()
        self.segment_group.setProperty("kind", "segmented")
        seg_layout = QHBoxLayout(self.segment_group)
        seg_layout.setContentsMargins(0, 0, 0, 0)
        seg_layout.setSpacing(0)
        self.pill_audio = StatusPill("Audio")
        self.pill_transcript = StatusPill("Transcription")
        self.pill_translate = StatusPill("Traduction")
        self.pill_summary = StatusPill("Résumé")
        self.pill_audio.setProperty("segment", "first")
        self.pill_transcript.setProperty("segment", "mid")
        self.pill_translate.setProperty("segment", "mid")
        self.pill_summary.setProperty("segment", "last")
        seg_layout.addWidget(self.pill_audio)
        seg_layout.addWidget(self.pill_transcript)
        seg_layout.addWidget(self.pill_translate)
        seg_layout.addWidget(self.pill_summary)
        right.addWidget(self.segment_group)

        self.btn_settings = QToolButton()
        self.btn_settings.setObjectName("SettingsButton")
        self.btn_settings.setProperty("variant", "ghost")
        self.btn_settings.setText("⚙️")
        right.addWidget(self.btn_settings)

        root.addLayout(right)

        self.btn_start.clicked.connect(self.start_clicked)
        self.btn_pause.clicked.connect(self.pause_clicked)
        self.btn_stop.clicked.connect(self.stop_clicked)

    def set_recording(self, recording: bool):
        """
        Centralize the recording UI state (colors + animation).
        """
        is_on = bool(recording)
        self.setProperty("recording", "true" if is_on else "false")
        self.lbl_rec.setProperty("recording", "true" if is_on else "false")
        self.style().polish(self)
        self.lbl_rec.style().polish(self.lbl_rec)
        try:
            if is_on:
                if self._rec_anim.state() != QPropertyAnimation.State.Running:
                    self._rec_anim.start()
            else:
                self._rec_anim.stop()
                self._rec_fx.setOpacity(1.0)
        except Exception:
            pass


class LiveTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("LiveTab")
        self.auto_scroll = True
        self.scroll_locked = False
        self._last_speaker = ""
        self._messages: list[LiveMessage] = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        top_actions = QHBoxLayout()
        top_actions.setSpacing(12)
        self.btn_lang_participants = QToolButton()
        self.btn_lang_participants.setObjectName("LiveOptionsButton")
        self.btn_lang_participants.setProperty("variant", "ghost")
        self.btn_lang_participants.setText("Auto")
        self.btn_lang_participants.setMinimumWidth(220)
        self.btn_lang_participants.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        self.btn_lang_participants.setToolTip("Langue du participant")
        self.btn_lang_participants.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)

        self.btn_lang_me = QToolButton()
        self.btn_lang_me.setObjectName("LiveOptionsButton")
        self.btn_lang_me.setProperty("variant", "ghost")
        self.btn_lang_me.setText("Auto")
        self.btn_lang_me.setMinimumWidth(220)
        self.btn_lang_me.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        self.btn_lang_me.setToolTip("Ma langue source")
        self.btn_lang_me.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)

        lang_part_wrap = QFrame()
        lang_part_wrap.setObjectName("LiveLangWrap")
        lang_part_layout = QVBoxLayout(lang_part_wrap)
        lang_part_layout.setContentsMargins(0, 0, 0, 0)
        lang_part_layout.setSpacing(4)
        lbl_part = QLabel("LANGUE DU PARTICIPANT")
        lbl_part.setObjectName("LiveOptionLabel")
        lang_part_layout.addWidget(lbl_part)
        lang_part_layout.addWidget(self.btn_lang_participants)

        arrow_lbl = QLabel("→")
        arrow_lbl.setObjectName("LiveArrow")
        arrow_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)

        lang_me_wrap = QFrame()
        lang_me_wrap.setObjectName("LiveLangWrap")
        lang_me_layout = QVBoxLayout(lang_me_wrap)
        lang_me_layout.setContentsMargins(0, 0, 0, 0)
        lang_me_layout.setSpacing(4)
        lbl_me = QLabel("MA LANGUE SOURCE")
        lbl_me.setObjectName("LiveOptionLabel")
        lang_me_layout.addWidget(lbl_me)
        lang_me_layout.addWidget(self.btn_lang_me)

        self.chk_autoscroll = QCheckBox("Auto-scroll")
        self.chk_autoscroll.setChecked(True)
        self.chk_live_enabled = QCheckBox("Live")
        self.chk_live_enabled.setChecked(True)
        self.btn_popout = QToolButton()
        self.btn_popout.setText("💬  Ouvrir chat")
        self.btn_popout.setObjectName("LiveActionPrimary")
        self.btn_popout.setProperty("variant", "primary")
        self.btn_rename_speaker = QToolButton()
        self.btn_rename_speaker.setText("✎  Renommer voix")
        self.btn_rename_speaker.setObjectName("LiveActionSecondary")
        self.btn_rename_speaker.setProperty("variant", "ghost")

        top_actions.addWidget(lang_part_wrap)
        top_actions.addWidget(arrow_lbl)
        top_actions.addWidget(lang_me_wrap)
        top_actions.addStretch(1)
        top_actions.addWidget(self.btn_popout)
        top_actions.addWidget(self.btn_rename_speaker)
        layout.addLayout(top_actions)

        self.preview_header = QFrame()
        self.preview_header.setObjectName("LivePreviewHeader")
        preview_layout = QHBoxLayout(self.preview_header)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(8)
        self.lbl_preview_title = QLabel("CHAT LIVE PREVIEW")
        self.lbl_preview_title.setObjectName("LivePreviewTitle")
        self.lbl_preview_title.setAlignment(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft)
        preview_layout.addWidget(self.lbl_preview_title)
        preview_layout.addStretch(1)
        self.btn_preview_history = QToolButton()
        self.btn_preview_history.setObjectName("LivePreviewHistory")
        self.btn_preview_history.setText("↺")
        self.btn_preview_history.setProperty("variant", "ghost")
        preview_layout.addWidget(self.btn_preview_history)
        layout.addWidget(self.preview_header)

        self.chat = LiveChatView()
        layout.addWidget(self.chat, 1)

        self.btn_scroll_bottom = QToolButton()
        self.btn_scroll_bottom.setText("↓ Revenir en bas")
        self.btn_scroll_bottom.setVisible(False)
        self.btn_scroll_bottom.setObjectName("ScrollBottomButton")
        self.btn_scroll_bottom.setProperty("variant", "ghost")

        bottom_actions = QHBoxLayout()
        bottom_actions.setSpacing(10)
        bottom_actions.addStretch(1)
        bottom_actions.addWidget(self.btn_scroll_bottom)
        layout.addLayout(bottom_actions)

        self.cmb_speaker_filter = QComboBox()
        self.cmb_speaker_filter.addItem("Tous")

        self.chat.scrolled.connect(self._on_scrolled)
        self.chat.resized.connect(lambda: None)
        self.btn_scroll_bottom.clicked.connect(self._scroll_to_bottom)
        self.chk_autoscroll.stateChanged.connect(self._toggle_autoscroll)
        self.chk_live_enabled.stateChanged.connect(self._on_live_toggle)

    def _toggle_autoscroll(self):
        self.auto_scroll = bool(self.chk_autoscroll.isChecked())
        if self.auto_scroll:
            self.scroll_locked = False
            self.btn_scroll_bottom.setVisible(False)

    def _on_scrolled(self):
        if not self.auto_scroll:
            return
        bar = self.chat.verticalScrollBar()
        if bar.value() < bar.maximum() - 6:
            self.scroll_locked = True
            self.btn_scroll_bottom.setVisible(True)

    def _scroll_to_bottom(self):
        self.chat.scrollToBottom()
        self.scroll_locked = False
        self.btn_scroll_bottom.setVisible(False)

    def _on_live_toggle(self):
        # handled by MainWindow
        return

    def add_message(self, msg: LiveMessage):
        self._messages.append(msg)
        should_scroll = bool(self.auto_scroll and not self.scroll_locked)
        self.chat.add_message(msg, keep_scroll=not should_scroll, scroll_to_bottom=should_scroll)

    def update_last_message(self, msg: LiveMessage):
        if not self._messages:
            self.add_message(msg)
            return
        self._messages[-1] = msg
        should_scroll = bool(self.auto_scroll and not self.scroll_locked)
        self.chat.update_last_message(msg, keep_scroll=not should_scroll, scroll_to_bottom=should_scroll)

    def get_last_message(self) -> Optional[LiveMessage]:
        return self._messages[-1] if self._messages else None

    def clear_messages(self):
        self._messages.clear()
        self._last_speaker = ""
        self.scroll_locked = False
        self.btn_scroll_bottom.setVisible(False)
        self.chat.clear_messages()


class SubjectWorker(QThread):
    finished_ok = pyqtSignal(str, str)  # session_dir, subject
    failed = pyqtSignal(str, str)

    def __init__(self, transcript_path: Path, session_dir: Path, cfg: dict, parent=None):
        super().__init__(parent)
        self.transcript_path = Path(transcript_path)
        self.session_dir = Path(session_dir)
        self.cfg = cfg or {}

    def run(self) -> None:
        try:
            from services.meeting_summary_service import generate_subject_from_transcript

            subject = generate_subject_from_transcript(self.transcript_path, self.cfg)
            if subject:
                try:
                    out_path = self.session_dir / "subject.txt"
                    out_path.write_text(subject.strip(), encoding="utf-8")
                except Exception:
                    pass
            self.finished_ok.emit(str(self.session_dir), subject or "")
        except Exception as e:
            self.failed.emit(str(self.session_dir), repr(e))


class StopOptionsDialog(QDialog):
    def __init__(self, cfg: dict, parent=None):
        super().__init__(parent)
        from PyQt6.QtWidgets import QSizePolicy as _QSizePolicy
        self.cfg = cfg or {}
        self.action = "cancel"
        self.setWindowTitle("Compte rendu")
        self.setMinimumWidth(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        title = QLabel("Avant traitement")
        title.setObjectName("DialogTitle")
        layout.addWidget(title)

        def add_section(text: str):
            lbl = QLabel(str(text).upper())
            lbl.setObjectName("SectionTitle")
            lbl.setSizePolicy(_QSizePolicy.Policy.Expanding, _QSizePolicy.Policy.Fixed)
            layout.addWidget(lbl)

        def add_card(text: str):
            frame = QFrame()
            frame.setProperty("kind", "card")
            v = QVBoxLayout(frame)
            v.setContentsMargins(12, 10, 12, 10)
            v.setSpacing(8)
            lbl = QLabel(str(text).upper())
            lbl.setObjectName("SectionTitle")
            lbl.setSizePolicy(_QSizePolicy.Policy.Expanding, _QSizePolicy.Policy.Fixed)
            v.addWidget(lbl)
            layout.addWidget(frame)
            return v

        card_main = add_card("Options principales")
        self.chk_transcript = QCheckBox("Générer la transcription (recommandé)")
        self.chk_transcript.setChecked(True)
        card_main.addWidget(self.chk_transcript)

        self.chk_report = QCheckBox("Générer le document DOCX")
        self.chk_report.setChecked(bool(self.cfg.get("postprocess_generate_docx", True)))
        card_main.addWidget(self.chk_report)

        self.chk_pplx = QCheckBox("Résumé structuré automatique (Perplexity)")
        self.chk_pplx.setChecked(bool(self.cfg.get("postprocess_enable_perplexity_summary", True)))
        card_main.addWidget(self.chk_pplx)

        card_tpl = add_card("Template")
        self.cmb_template = QComboBox()
        self.cmb_template.addItems(
            ["Compte rendu pro", "Recrutement", "Vidéo YouTube", "Webinaire", "Réunion Discord (asso)"]
        )
        cur_tpl = str(self.cfg.get("summary_template") or "Compte rendu pro")
        idx = self.cmb_template.findText(cur_tpl)
        if idx >= 0:
            self.cmb_template.setCurrentIndex(idx)
        card_tpl.addWidget(QLabel("Template du compte rendu"))
        card_tpl.addWidget(self.cmb_template)

        card_rules = add_card("Règles & qualité")

        self.cmb_lang = QComboBox()
        self.cmb_lang.addItem("Automatique (mix)", "auto")
        self.cmb_lang.addItem("Francais", "fr")
        self.cmb_lang.addItem("Anglais", "en")
        cur_lang = (self.cfg.get("postprocess_language") or "auto").lower()
        self.cmb_lang.setCurrentIndex(0 if cur_lang == "auto" else (1 if cur_lang == "fr" else 2))
        card_rules.addWidget(QLabel("Langue du compte rendu"))
        card_rules.addWidget(self.cmb_lang)

        self.cmb_quality = QComboBox()
        self.cmb_quality.addItem("Standard (plus rapide)", "standard")
        self.cmb_quality.addItem("Précis (meilleur, plus lent)", "precise")
        cur_q = (self.cfg.get("postprocess_quality") or "standard").lower()
        self.cmb_quality.setCurrentIndex(0 if cur_q == "standard" else 1)
        card_rules.addWidget(QLabel("Qualité du compte rendu"))
        card_rules.addWidget(self.cmb_quality)

        self.cmb_diar_mode = QComboBox()
        self.cmb_diar_mode.addItem("Par voix (plus précis, nécessite token HF)", "voice")
        self.cmb_diar_mode.addItem("Par source (simple: micro/participants)", "source")
        cur_mode = str(self.cfg.get("postprocess_diarization_mode") or "").lower()
        if not cur_mode:
            if not bool(self.cfg.get("postprocess_enable_diarization", self.cfg.get("enable_diarization", True))):
                cur_mode = "source"
            else:
                cur_mode = "voice"
        self.cmb_diar_mode.setCurrentIndex(0 if cur_mode == "voice" else 1)
        card_people = add_card("Personnes & voix")
        card_people.addWidget(QLabel("Séparation des voix"))
        card_people.addWidget(self.cmb_diar_mode)

        self.chk_extract = QCheckBox("Déduire les noms des participants (IA)")
        self.chk_extract.setChecked(bool(self.cfg.get("postprocess_extract_participants", False)))
        card_people.addWidget(self.chk_extract)

        card_perf = add_card("Performance")
        self.cmb_device = QComboBox()
        self.cmb_device.addItem("Auto (GPU si dispo)", "auto")
        self.cmb_device.addItem("GPU (CUDA)", "cuda")
        self.cmb_device.addItem("CPU", "cpu")
        cur_dev = str(self.cfg.get("postprocess_device") or self.cfg.get("device") or "auto").lower()
        self.cmb_device.setCurrentIndex(0 if cur_dev == "auto" else (1 if cur_dev == "cuda" else 2))
        card_perf.addWidget(QLabel("Exécution"))
        card_perf.addWidget(self.cmb_device)

        btns = QHBoxLayout()
        btns.addStretch(1)
        self.btn_delete = PulseButton("Supprimer l'enregistrement")
        self.btn_delete.setObjectName("BtnDelete")
        self.btn_delete.setProperty("variant", "solid")
        self.btn_delete.setProperty("tone", "danger")
        self.btn_cancel = PulseButton("Annuler")
        self.btn_cancel.setObjectName("BtnCancel")
        self.btn_cancel.setProperty("variant", "solid")
        self.btn_cancel.setProperty("tone", "neutral")
        self.btn_start = PulseButton("Commencer")
        self.btn_start.setObjectName("BtnStart")
        self.btn_start.setProperty("variant", "solid")
        self.btn_start.setProperty("tone", "success")
        btns.addWidget(self.btn_delete)
        btns.addWidget(self.btn_cancel)
        btns.addWidget(self.btn_start)
        layout.addLayout(btns)

        # Uniform button sizes for consistent UI
        widths = [self.btn_delete.sizeHint().width(), self.btn_cancel.sizeHint().width(), self.btn_start.sizeHint().width()]
        target_w = max(widths) + 6
        for b in (self.btn_delete, self.btn_cancel, self.btn_start):
            b.setFixedWidth(target_w)

        self.btn_start.clicked.connect(self._on_start)
        self.btn_cancel.clicked.connect(self._on_cancel)
        self.btn_delete.clicked.connect(self._on_delete)
        self.chk_transcript.stateChanged.connect(self._apply_enabled_states)
        self.chk_report.stateChanged.connect(self._apply_enabled_states)
        self._apply_enabled_states()
        self._apply_diarization_availability()

    def _apply_diarization_availability(self):
        token = getsecret(self.cfg, "hf_token") or ""
        if not token:
            if self.cmb_diar_mode.currentData() == "voice":
                self.cmb_diar_mode.setCurrentIndex(1)

    def _apply_enabled_states(self):
        transcript_on = bool(self.chk_transcript.isChecked())
        self.chk_report.setEnabled(transcript_on)
        self.cmb_lang.setEnabled(transcript_on)
        self.cmb_quality.setEnabled(transcript_on)
        self.cmb_diar_mode.setEnabled(transcript_on)
        report_on = transcript_on and bool(self.chk_report.isChecked())
        self.chk_pplx.setEnabled(report_on)
        self.cmb_template.setEnabled(report_on)
        self.chk_extract.setEnabled(report_on)

    def _on_start(self):
        self.action = "start"
        self.accept()

    def _on_cancel(self):
        self.action = "cancel"
        self.reject()

    def _on_delete(self):
        self.action = "delete"
        self.accept()

    def get_options(self) -> dict:
        return {
            "generate_transcript": bool(self.chk_transcript.isChecked()),
            "generate_report": bool(self.chk_report.isChecked()),
            "postprocess_language": str(self.cmb_lang.currentData()),
            "postprocess_quality": str(self.cmb_quality.currentData()),
            "postprocess_generate_docx": bool(self.chk_report.isChecked()),
            "postprocess_enable_perplexity_summary": bool(self.chk_pplx.isChecked()),
            "postprocess_diarization_mode": str(self.cmb_diar_mode.currentData()),
            "postprocess_enable_diarization": bool(self.cmb_diar_mode.currentData() == "voice"),
            "postprocess_extract_participants": bool(self.chk_extract.isChecked()),
            "summary_template": str(self.cmb_template.currentText()),
            "postprocess_device": str(self.cmb_device.currentData()),
        }


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_NAME)
        self.setMinimumSize(980, 620)

        self.cfg = load_config()

        self.live_queue = queue.Queue(maxsize=300)
        self.recorder = self._build_recorder_from_cfg()

        self._configure_live_source_queue()

        self.session_dir: Optional[Path] = None
        self.live_thread = None
        self.live_chat_window: Optional[LiveChatWindow] = None
        self.pp_thread = None
        self._last_live_message: Optional[LiveMessage] = None
        self._paused = False
        self._current_session: Optional[Path] = None
        self._transcript_segments = []
        self._filtered_segments = []
        self._transcript_tags = []
        self._current_docx: Optional[Path] = None
        self._history_signals_connected = False
        self._summary_worker: Optional[SummaryWorker] = None
        self._summary_open_after = False
        self._debug_timer = QTimer(self)
        self._debug_timer.setInterval(1000)
        self._debug_timer.timeout.connect(self._poll_debug_log)
        self._history_loading = False
        self._subject_queue: list[Path] = []
        self._subject_worker: Optional[SubjectWorker] = None
        self._live_start_time: Optional[float] = None
        self._live_messages: list[LiveMessage] = []
        self._live_speaker_aliases: dict[str, str] = dict(self.cfg.get("live_speaker_aliases") or {})
        self._chunk_enabled = bool(self.cfg.get("postprocess_chunked", True))
        self._recording_action_lock = False
        self._retained_threads: list[QThread] = []
        self._chunk_event_queue = queue.Queue()
        self._chunk_poll_timer = QTimer(self)
        self._chunk_poll_timer.setInterval(300)
        self._chunk_poll_timer.timeout.connect(self._poll_chunk_events)
        self._reset_chunk_state(init=True)

        self._build_ui()
        self._apply_cfg_to_ui()

        self.timer = QTimer(self)
        self.timer.timeout.connect(self._tick_timer)
        self._seconds = 0

        self.pp_progress_timer = QTimer(self)
        self.pp_progress_timer.setInterval(500)
        self.pp_progress_timer.timeout.connect(self._poll_postprocess_progress)
        self._pp_last_payload = None
        self._pp_last_payload_ts = 0.0
        self._pp_display_percent = 0
        self.price_timer = QTimer(self)
        self.price_timer.setInterval(30000)
        self.price_timer.timeout.connect(self._update_price_estimate)

    def _parse_cfg_device_id(self, key: str) -> Optional[int]:
        raw = self.cfg.get(key, None)
        if raw is None:
            return None
        try:
            val = int(raw)
        except Exception:
            return None
        return val if val >= 0 else None

    def _non_windows_audio_defaults(self) -> Optional[tuple[int, int]]:
        """
        On macOS/Linux, both "participants" and "micro" are input devices.
        Participants should ideally be a virtual loopback input (BlackHole/Loopback/Soundflower).
        """
        if sd is None:
            return None
        try:
            devices = sd.query_devices()
        except Exception:
            return None

        inputs: list[tuple[int, str]] = []
        for idx, dev in enumerate(devices):
            try:
                if int(dev.get("max_input_channels", 0) or 0) > 0:
                    inputs.append((int(idx), str(dev.get("name", ""))))
            except Exception:
                continue

        if not inputs:
            return None

        valid_ids = {idx for idx, _ in inputs}
        default_in = None
        try:
            dflt = sd.default.device
            if isinstance(dflt, (list, tuple)):
                default_in = int(dflt[0]) if dflt else None
            else:
                default_in = int(dflt)
        except Exception:
            default_in = None

        if default_in not in valid_ids:
            default_in = inputs[0][0]

        # Prefer known virtual loopback names for participants.
        loopback_words = (
            "blackhole",
            "loopback",
            "soundflower",
            "vb-cable",
            "stereo mix",
            "monitor",
            "virtual",
        )
        participants_id = None
        for idx, name in inputs:
            low = name.lower()
            if any(k in low for k in loopback_words):
                participants_id = idx
                break

        if participants_id is None:
            # Fallback to another input if available, else same as mic.
            participants_id = default_in
            for idx, _name in inputs:
                if idx != default_in:
                    participants_id = idx
                    break

        return int(participants_id), int(default_in)

    def _resolve_audio_devices_from_cfg(self) -> tuple[int, int]:
        part_id = self._parse_cfg_device_id("participantsoutputdeviceid")
        mic_id = self._parse_cfg_device_id("microdeviceid")

        if part_id is not None and mic_id is not None:
            return int(part_id), int(mic_id)

        if os.name == "nt":
            raise RuntimeError(
                "Configuration audio manquante : ouvre Configuration et sélectionne "
                "la sortie Windows (participants) et le micro."
            )

        detected = self._non_windows_audio_defaults()
        if not detected:
            raise RuntimeError(
                "Aucune source audio d'entrée détectée. Sur macOS, installe BlackHole/Loopback "
                "puis configure 'Sortie audio' et 'Entrée audio' dans Configuration."
            )

        part_id, mic_id = detected
        self.cfg["participantsoutputdeviceid"] = int(part_id)
        self.cfg["microdeviceid"] = int(mic_id)
        save_config(self.cfg)
        return int(part_id), int(mic_id)

    def _build_recorder_from_cfg(self) -> RecorderService:
        part_id, mic_id = self._resolve_audio_devices_from_cfg()
        configured_dir = self.cfg.get("sessions_dir") or DEFAULT_SESSIONS_DIR
        sessions_dir = self._resolve_sessions_dir(configured_dir)
        return RecorderService(
            participants_output_device_id=int(part_id),
            mic_device_id=int(mic_id),
            output_root=sessions_dir,
        )

    def _default_sessions_dir(self) -> Path:
        docs = Path.home() / "Documents"
        base = docs if docs.exists() else Path.home()
        return base / APP_NAME / "recordings"

    def _is_writable_dir(self, path: Path) -> bool:
        try:
            path.mkdir(parents=True, exist_ok=True)
            probe = path / ".write_test.tmp"
            probe.write_text("ok", encoding="utf-8")
            probe.unlink(missing_ok=True)
            return True
        except Exception:
            return False

    def _resolve_sessions_dir(self, configured_dir) -> Path:
        configured = Path(str(configured_dir)) if configured_dir else self._default_sessions_dir()
        fallback = self._default_sessions_dir()
        for candidate in (configured, fallback):
            if self._is_writable_dir(candidate):
                final_dir = candidate
                break
        else:
            raise RuntimeError(
                "Impossible d'écrire dans le dossier des sessions. "
                "Choisis un dossier utilisateur dans la configuration."
            )

        final_str = str(final_dir)
        if str(self.cfg.get("sessions_dir") or "") != final_str:
            self.cfg["sessions_dir"] = final_str
            save_config(self.cfg)
            if str(configured) != final_str:
                log_line(f"[SESSIONS] Fallback dossier sessions: {final_str}")
        return final_dir

    def _reload_recorder_from_cfg(self) -> None:
        if bool(getattr(getattr(self, "recorder", None), "_running", False)):
            return
        self.recorder = self._build_recorder_from_cfg()
        self._configure_live_source_queue()


    def _build_ui(self):
        root = QWidget()
        layout = QVBoxLayout(root)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.topbar = TopBar()
        layout.addWidget(self.topbar)

        # Visual breathing room between session controls and navigation/content.
        gap = QFrame()
        gap.setObjectName("SectionGap")
        gap.setFixedHeight(10)
        layout.addWidget(gap)

        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)
        try:
            # Workaround: ensure tab visuals update immediately on click (some styles/QSS combos can lag).
            self.tabs.currentChanged.connect(lambda _i: self.tabs.tabBar().update())
        except Exception:
            pass
        tabs_wrap = QFrame()
        tabs_wrap.setObjectName("TabsWrap")
        tabs_layout = QVBoxLayout(tabs_wrap)
        tabs_layout.setContentsMargins(12, 0, 12, 0)
        tabs_layout.addWidget(self.tabs)
        layout.addWidget(tabs_wrap, 1)

        self.live_tab = LiveTab()
        self.tabs.addTab(self.live_tab, "LIVE")

        self.transcription_tab = self._build_transcription_tab()
        self.tabs.addTab(self.transcription_tab, "TRANSCRIPTION")

        self.summary_tab = self._build_summary_tab()
        self.tabs.addTab(self.summary_tab, "RÉSUMÉ")

        self.history_tab = self._build_history_tab()
        self.tabs.addTab(self.history_tab, "HISTORIQUE")

        self.pp_progress = QProgressBar()
        self.pp_progress.setRange(0, 100)
        self.pp_progress.setValue(0)
        self.pp_progress.setVisible(False)
        self.lbl_progress = QLabel("")
        self.lbl_progress.setObjectName("ProgressLabel")
        self.lbl_progress.setVisible(False)
        progress_row = QHBoxLayout()
        progress_row.setContentsMargins(12, 0, 12, 0)
        progress_row.setSpacing(8)
        progress_row.addWidget(self.lbl_progress, 1)
        self.btn_cancel_pp = QPushButton("Annuler")
        self.btn_cancel_pp.setObjectName("BtnCancel")
        self.btn_cancel_pp.setProperty("variant", "ghost")
        self.btn_cancel_pp.setVisible(False)
        progress_row.addWidget(self.btn_cancel_pp)
        layout.addLayout(progress_row)
        layout.addWidget(self.pp_progress)

        self.debug_panel = QPlainTextEdit()
        self.debug_panel.setObjectName("DebugPanel")
        self.debug_panel.setReadOnly(True)
        self.debug_panel.setVisible(False)
        self.debug_panel.setMaximumHeight(160)
        layout.addWidget(self.debug_panel)

        bottom = QFrame()
        bottom.setObjectName("BottomBar")
        bottom_l = QHBoxLayout(bottom)
        bottom_l.setContentsMargins(12, 6, 12, 6)
        bottom_l.setSpacing(8)
        bottom_l.addWidget(self.live_tab.chk_autoscroll)
        bottom_l.addWidget(self.live_tab.chk_live_enabled)
        sep = QFrame()
        sep.setObjectName("BottomDivider")
        sep.setFixedSize(1, 16)
        bottom_l.addWidget(sep)
        self.lbl_price = QLabel("Prix transcription: OpenAI 0,00 $ | Deepgram 0,00 $")
        self.lbl_price.setObjectName("PriceLabel")
        self.lbl_price.setToolTip("Tarifs indicatifs en USD (mise à jour toutes les 30s).")
        bottom_l.addWidget(self.lbl_price)
        bottom_l.addStretch(1)
        self.lbl_cost = QLabel("Coût: 0 tokens")
        self.lbl_cost.setObjectName("CostLabelBottom")
        bottom_l.addWidget(self.lbl_cost)
        layout.addWidget(bottom)

        self.setCentralWidget(root)

        self.topbar.start_clicked.connect(self._on_start)
        self.topbar.stop_clicked.connect(self._on_stop)
        self.topbar.pause_clicked.connect(self._on_pause)
        self.topbar.btn_reset.clicked.connect(self._on_reset)
        self.topbar.btn_settings.clicked.connect(self._on_settings)

        self._init_live_menus()

        self._load_history()

        self.live_tab.chk_live_enabled.stateChanged.connect(self._on_live_enabled_changed)
        self.live_tab.btn_popout.clicked.connect(self._on_open_live_chat_window)
        self.live_tab.btn_rename_speaker.clicked.connect(self._on_rename_live_speaker)

    def _build_transcription_tab(self) -> QWidget:
        root = QWidget()
        layout = QVBoxLayout(root)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        tools = QHBoxLayout()
        self.search_transcript = QLineEdit()
        self.search_transcript.setPlaceholderText("Rechercher...")
        self.btn_clear_search = QPushButton("Clear")
        self.btn_clear_search.setProperty("variant", "ghost")
        self.cmb_filter_speaker = QComboBox()
        self.cmb_filter_speaker.addItem("Tous")
        self.btn_clean = QPushButton("Nettoyage auto")
        self.btn_clean.setProperty("variant", "ghost")
        tools.addWidget(self.search_transcript, 1)
        tools.addWidget(self.btn_clear_search)
        tools.addWidget(self.cmb_filter_speaker)
        tools.addWidget(self.btn_clean)
        layout.addLayout(tools)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        self.list_transcript = QListWidget()
        self.list_transcript.setObjectName("TranscriptList")
        self.list_transcript.setVerticalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        splitter.addWidget(self.list_transcript)

        detail = QWidget()
        dl = QVBoxLayout(detail)
        dl.setContentsMargins(8, 8, 8, 8)
        self.txt_detail = QPlainTextEdit()
        self.txt_detail.setPlaceholderText("Détails du segment...")
        self.txt_detail.setReadOnly(False)
        dl.addWidget(QLabel("Original / Traduction"))
        dl.addWidget(self.txt_detail, 1)
        chip_row = QHBoxLayout()
        self.chk_tag_important = QCheckBox("Important")
        self.chk_tag_action = QCheckBox("Action")
        self.chk_tag_decision = QCheckBox("Décision")
        chip_row.addWidget(self.chk_tag_important)
        chip_row.addWidget(self.chk_tag_action)
        chip_row.addWidget(self.chk_tag_decision)
        chip_row.addStretch(1)
        dl.addLayout(chip_row)
        splitter.addWidget(detail)

        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 2)
        layout.addWidget(splitter, 1)
        return root

    def _build_summary_tab(self) -> QWidget:
        root = QWidget()
        layout = QVBoxLayout(root)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        tools = QHBoxLayout()
        self.cmb_template = QComboBox()
        self.cmb_template.addItems(
            [
                "Compte rendu pro",
                "Recrutement",
                "Vidéo YouTube",
                "Webinaire",
                "Réunion Discord (asso)",
            ]
        )
        cur_tpl = str(self.cfg.get("summary_template") or "Compte rendu pro")
        idx = self.cmb_template.findText(cur_tpl)
        if idx >= 0:
            self.cmb_template.setCurrentIndex(idx)
        self.btn_generate = QPushButton("Générer")
        self.btn_generate.setProperty("variant", "primary")
        self.btn_regenerate = QPushButton("Régénérer")
        self.btn_regenerate.setProperty("variant", "ghost")
        self.btn_copy_summary = QPushButton("Copier")
        self.btn_copy_summary.setProperty("variant", "ghost")
        self.btn_export_docx = QPushButton("Exporter DOCX")
        self.btn_export_docx.setProperty("variant", "primary")
        self.btn_open_docx = QPushButton("Ouvrir DOCX")
        self.btn_open_docx.setProperty("variant", "ghost")
        self.btn_open_docx.setEnabled(False)
        tools.addWidget(QLabel("Template"))
        tools.addWidget(self.cmb_template)
        tools.addStretch(1)
        tools.addWidget(self.btn_generate)
        tools.addWidget(self.btn_regenerate)
        tools.addWidget(self.btn_copy_summary)
        tools.addWidget(self.btn_export_docx)
        tools.addWidget(self.btn_open_docx)
        layout.addLayout(tools)

        self.txt_summary = QPlainTextEdit()
        self.txt_summary.setPlaceholderText("Résumé de réunion...")
        layout.addWidget(self.txt_summary, 1)
        return root

    def _build_history_tab(self) -> QWidget:
        root = QWidget()
        layout = QVBoxLayout(root)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        self.table_history = QTableWidget(0, 7)
        self.table_history.setHorizontalHeaderLabels(
            ["Sujet", "Date", "Nom session", "Durée", "Nb speakers", "Exports", "Action"]
        )
        self.table_history.horizontalHeader().setStretchLastSection(True)
        self.table_history.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table_history.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked)
        self.table_history.setObjectName("HistoryTable")
        layout.addWidget(self.table_history, 1)
        return root

    def _load_history(self, focus_session: Optional[Path] = None):
        self._history_loading = True
        self.table_history.setRowCount(0)
        sessions_dir = Path(self.cfg.get("sessions_dir") or DEFAULT_SESSIONS_DIR)
        if not sessions_dir.exists():
            self._history_loading = False
            return

        sessions = []
        for date_dir in sessions_dir.iterdir():
            if not date_dir.is_dir():
                continue
            for session_dir in date_dir.iterdir():
                if session_dir.is_dir():
                    sessions.append(session_dir)

        sessions.sort(key=lambda p: p.stat().st_mtime, reverse=True)

        self._subject_queue = []
        for session_dir in sessions:
            info = self._session_info(session_dir)
            row = self.table_history.rowCount()
            self.table_history.insertRow(row)
            subject = info.get("subject", "—")
            self._set_history_item(row, 0, subject or "—", str(session_dir), editable=True)
            self._set_history_item(row, 1, info.get("date", ""))
            self._set_history_item(row, 2, info.get("name", session_dir.name))
            self._set_history_item(row, 3, info.get("duration", "—"))
            self._set_history_item(row, 4, info.get("speakers", "—"))
            self._set_history_item(row, 5, info.get("exports", "—"))
            btn = QPushButton("Ouvrir dossier")
            btn.setProperty("variant", "ghost")
            btn.clicked.connect(lambda _checked=False, p=session_dir: self._open_folder(p))
            self.table_history.setCellWidget(row, 6, btn)
            if subject in ("", "—") and info.get("can_autofill_subject"):
                self._subject_queue.append(session_dir)

        if self.table_history.rowCount() > 0:
            target_row = 0
            if focus_session:
                for r in range(self.table_history.rowCount()):
                    item = self.table_history.item(r, 0)
                    if item and item.data(Qt.ItemDataRole.UserRole) == str(focus_session):
                        target_row = r
                        break
            self.table_history.selectRow(target_row)
            self._load_session_from_row(target_row)
        else:
            self.list_transcript.clear()
            self.txt_summary.setPlainText("Aucune session enregistrée.")
            self.cmb_filter_speaker.clear()
            self.cmb_filter_speaker.addItem("Tous")
            self.btn_open_docx.setEnabled(False)

        if not self._history_signals_connected:
            self.table_history.itemSelectionChanged.connect(self._on_history_selection)
            self.table_history.itemChanged.connect(self._on_history_item_changed)
            self.btn_open_docx.clicked.connect(self._open_current_docx)
            self.btn_clear_search.clicked.connect(self._on_clear_search)
            self.search_transcript.textChanged.connect(self._apply_transcript_filter)
            self.cmb_filter_speaker.currentTextChanged.connect(self._apply_transcript_filter)
            self.list_transcript.currentRowChanged.connect(self._on_transcript_selected)
            self.btn_generate.clicked.connect(self._on_generate_summary)
            self.btn_regenerate.clicked.connect(self._on_regenerate_summary)
            self.btn_copy_summary.clicked.connect(self._on_copy_summary)
            self.btn_export_docx.clicked.connect(self._on_export_docx)
            self.btn_clean.clicked.connect(self._on_clean_transcript)
            self.chk_tag_important.stateChanged.connect(lambda: self._on_tag_changed("important"))
            self.chk_tag_action.stateChanged.connect(lambda: self._on_tag_changed("action"))
            self.chk_tag_decision.stateChanged.connect(lambda: self._on_tag_changed("decision"))
            self.cmb_template.currentTextChanged.connect(self._on_template_changed)
            self.btn_cancel_pp.clicked.connect(self._on_cancel_postprocess)
            self._history_signals_connected = True

        self._history_loading = False
        self._start_subject_queue()

    def _set_history_item(self, row: int, col: int, text: str, data=None, editable: bool = False):
        item = QTableWidgetItem(text)
        if data is not None:
            item.setData(Qt.ItemDataRole.UserRole, data)
        flags = item.flags()
        if editable:
            item.setFlags(flags | Qt.ItemFlag.ItemIsEditable)
        else:
            item.setFlags(flags & ~Qt.ItemFlag.ItemIsEditable)
        self.table_history.setItem(row, col, item)

    def _on_history_selection(self):
        sel = self.table_history.selectionModel().selectedRows()
        if not sel:
            return
        row = sel[0].row()
        self._load_session_from_row(row)

    def _load_session_from_row(self, row: int):
        item = self.table_history.item(row, 0)
        if not item:
            return
        session_dir = self._session_dir_from_row(row)
        if session_dir:
            self._load_session(session_dir)

    def _session_dir_from_row(self, row: int) -> Optional[Path]:
        item = self.table_history.item(row, 0)
        if not item:
            return None
        data = item.data(Qt.ItemDataRole.UserRole)
        if not data:
            return None
        p = Path(str(data))
        return p if p.exists() else None

    def _load_session(self, session_dir: Path):
        if not session_dir or not session_dir.exists():
            return
        self._current_session = session_dir
        self._load_transcript(session_dir)
        self._load_summary(session_dir)

    def _session_info(self, session_dir: Path) -> dict:
        info = {"date": session_dir.parent.name, "name": session_dir.name}
        args_path = session_dir / "postprocess_args.json"
        if args_path.exists():
            try:
                data = json.loads(args_path.read_text(encoding="utf-8"))
                cfg = data.get("cfg") or {}
                name = cfg.get("session_name") or cfg.get("name")
                if name:
                    info["name"] = str(name)
            except Exception:
                pass

        info["duration"] = self._calc_session_duration(session_dir)
        info["speakers"] = self._calc_speaker_count(session_dir)
        info["exports"] = "DOCX" if self._find_docx(session_dir) else "—"
        info["subject"] = self._load_subject(session_dir)
        info["can_autofill_subject"] = self._can_autofill_subject()
        return info

    def _calc_session_duration(self, session_dir: Path) -> str:
        wavs = list(session_dir.glob("*.wav"))
        if not wavs:
            return "—"
        target_wavs = [p for p in wavs if "participants" in p.name.lower()]
        if not target_wavs:
            target_wavs = wavs
        seconds = 0
        for wav_path in sorted(target_wavs):
            try:
                with wave.open(str(wav_path), "rb") as wf:
                    frames = wf.getnframes()
                    rate = wf.getframerate() or 1
                    seconds += int(frames / rate)
            except Exception:
                continue
        if seconds <= 0:
            return "—"
        hh = seconds // 3600
        mm = (seconds % 3600) // 60
        ss = seconds % 60
        return f"{hh:02d}:{mm:02d}:{ss:02d}"

    def _calc_speaker_count(self, session_dir: Path) -> str:
        tpath = self._find_transcript_path(session_dir)
        if not tpath:
            return "—"
        speakers = set()
        try:
            for line in tpath.read_text(encoding="utf-8").splitlines():
                m = re.match(r"^\[[^\]]+\]\s*([^:]+):", line.strip())
                if m:
                    speakers.add(m.group(1).strip())
        except Exception:
            return "—"
        return str(len(speakers)) if speakers else "—"

    def _find_transcript_path(self, session_dir: Path) -> Optional[Path]:
        preferred = [
            session_dir / "transcript_speakers_fr.txt",
            session_dir / "transcript-speakers.mix.txt",
        ]
        for p in preferred:
            if p.exists():
                return p
        for p in session_dir.glob("transcript*.txt"):
            return p
        return None

    def _load_subject(self, session_dir: Path) -> str:
        p = session_dir / "subject.txt"
        if p.exists():
            try:
                return p.read_text(encoding="utf-8").strip()
            except Exception:
                return ""
        return ""

    def _can_autofill_subject(self) -> bool:
        if not bool(self.cfg.get("history_auto_subject", True)):
            return False
        from services.meeting_summary_service import _get_perplexity_key

        key, _ = _get_perplexity_key(self.cfg or {})
        return bool(key)

    def _parse_transcript_segments(self, path: Path) -> list[dict]:
        segments = []
        if not path or not path.exists():
            return segments
        pattern = re.compile(r"^\[(?P<start>\d{2}:\d{2}:\d{2})\s*-\s*\d{2}:\d{2}:\d{2}\]\s*(?P<spk>[^:]+):\s*(?P<txt>.*)$")
        pattern_no_spk = re.compile(r"^\[(?P<start>\d{2}:\d{2}:\d{2})\s*-\s*\d{2}:\d{2}:\d{2}\]\s*(?P<txt>.*)$")
        for idx, line in enumerate(path.read_text(encoding="utf-8").splitlines()):
            line = line.strip()
            if not line:
                continue
            m = pattern.match(line)
            if m:
                segments.append(
                    {
                        "time": m.group("start"),
                        "speaker": m.group("spk").strip(),
                        "text": m.group("txt").strip(),
                        "idx": idx,
                    }
                )
            else:
                m2 = pattern_no_spk.match(line)
                if m2:
                    segments.append(
                        {
                            "time": m2.group("start"),
                            "speaker": "",
                            "text": m2.group("txt").strip(),
                            "idx": idx,
                        }
                    )
                else:
                    segments.append({"time": "", "speaker": "", "text": line, "idx": idx})
        return segments

    def _load_transcript(self, session_dir: Path):
        self._transcript_segments = []
        self._filtered_segments = []
        self._transcript_tags = []
        self.list_transcript.clear()
        self.cmb_filter_speaker.clear()
        self.cmb_filter_speaker.addItem("Tous")
        self.txt_detail.clear()

        tpath = self._find_transcript_path(session_dir)
        if not tpath:
            self.list_transcript.addItem("Aucune transcription disponible.")
            return
        self._transcript_segments = self._parse_transcript_segments(tpath)
        self._transcript_tags = [set() for _ in range(len(self._transcript_segments))]
        speakers = sorted({seg["speaker"] for seg in self._transcript_segments if seg.get("speaker")})
        for s in speakers:
            self.cmb_filter_speaker.addItem(s)
        self._apply_transcript_filter()

    def _apply_transcript_filter(self):
        if not hasattr(self, "list_transcript"):
            return
        query = (self.search_transcript.text() or "").strip().lower()
        speaker = self.cmb_filter_speaker.currentText().strip()
        self.list_transcript.clear()
        self._filtered_segments = []
        for seg in self._transcript_segments:
            spk = seg.get("speaker", "")
            txt = seg.get("text", "")
            if speaker and speaker != "Tous" and spk != speaker:
                continue
            if query and query not in txt.lower():
                continue
            label = self._format_transcript_item(seg)
            self.list_transcript.addItem(label)
            self._filtered_segments.append(seg)

    def _on_transcript_selected(self, row: int):
        if row < 0 or row >= len(self._filtered_segments):
            return
        seg = self._filtered_segments[row]
        header = ""
        if seg.get("time") or seg.get("speaker"):
            header = f"{seg.get('time','')} | {seg.get('speaker','')}".strip()
        body = seg.get("text", "")
        self.txt_detail.setPlainText((header + "\n" + body).strip())
        tags = self._get_tags_for_segment(seg)
        self.chk_tag_important.blockSignals(True)
        self.chk_tag_action.blockSignals(True)
        self.chk_tag_decision.blockSignals(True)
        self.chk_tag_important.setChecked("important" in tags)
        self.chk_tag_action.setChecked("action" in tags)
        self.chk_tag_decision.setChecked("decision" in tags)
        self.chk_tag_important.blockSignals(False)
        self.chk_tag_action.blockSignals(False)
        self.chk_tag_decision.blockSignals(False)

    def _on_clear_search(self):
        self.search_transcript.clear()
        self._apply_transcript_filter()

    def _format_transcript_item(self, seg: dict) -> str:
        tags = self._get_tags_for_segment(seg)
        tag_txt = ""
        if tags:
            parts = []
            if "important" in tags:
                parts.append("★")
            if "action" in tags:
                parts.append("Action")
            if "decision" in tags:
                parts.append("Décision")
            tag_txt = " [" + ", ".join(parts) + "]"
        return f"{seg.get('time','')} | {seg.get('speaker','')}: {seg.get('text','')}{tag_txt}".strip()

    def _get_tags_for_segment(self, seg: dict) -> set:
        idx = seg.get("idx")
        if idx is None or idx < 0 or idx >= len(self._transcript_tags):
            return set()
        return self._transcript_tags[idx]

    def _on_tag_changed(self, tag: str):
        row = self.list_transcript.currentRow()
        if row < 0 or row >= len(self._filtered_segments):
            return
        seg = self._filtered_segments[row]
        tags = self._get_tags_for_segment(seg)
        chk_map = {
            "important": self.chk_tag_important,
            "action": self.chk_tag_action,
            "decision": self.chk_tag_decision,
        }
        if tag in chk_map and chk_map[tag].isChecked():
            tags.add(tag)
        else:
            tags.discard(tag)
        item = self.list_transcript.item(row)
        if item:
            item.setText(self._format_transcript_item(seg))

    def _clean_text(self, text: str) -> str:
        t = (text or "").strip()
        t = re.sub(r"\s+", " ", t)
        t = re.sub(r"\s+([,.;:!?])", r"\1", t)
        if t and t[0].islower():
            t = t[0].upper() + t[1:]
        return t

    def _on_clean_transcript(self):
        if not self._transcript_segments:
            return
        for seg in self._transcript_segments:
            seg["text"] = self._clean_text(seg.get("text", ""))
        self._apply_transcript_filter()
        QMessageBox.information(self, "Nettoyage auto", "Nettoyage appliqué à l'affichage.")

    def _find_docx(self, session_dir: Path) -> Optional[Path]:
        candidates = [
            session_dir / "Résumé de Réunion.docx",
            session_dir / "Resume de Reunion.docx",
        ]
        for p in candidates:
            if p.exists():
                return p
        for p in session_dir.glob("*.docx"):
            return p
        return None

    def _load_summary(self, session_dir: Path):
        self._current_docx = None
        summary_text = ""

        txt_candidates = [
            session_dir / "summary.txt",
            session_dir / "summary_fr.txt",
            session_dir / "resume.txt",
            session_dir / "resume_fr.txt",
        ]
        for p in txt_candidates:
            if p.exists():
                try:
                    summary_text = p.read_text(encoding="utf-8")
                    break
                except Exception:
                    pass

        docx_path = self._find_docx(session_dir)
        self._current_docx = docx_path
        if not summary_text and docx_path:
            try:
                from docx import Document

                doc = Document(str(docx_path))
                summary_text = "\n".join(p.text for p in doc.paragraphs if p.text).strip()
            except Exception:
                summary_text = "Résumé disponible dans le fichier DOCX (ouvrir pour voir)."

        if summary_text:
            self.txt_summary.setPlainText(summary_text)
        else:
            self.txt_summary.setPlainText("Aucun résumé disponible pour cette session.")

        self.btn_open_docx.setEnabled(bool(docx_path))

    def _open_current_docx(self):
        if not self._current_docx:
            return
        if not self._open_path(self._current_docx):
            QMessageBox.warning(self, "DOCX", "Impossible d'ouvrir le fichier DOCX.")

    def _open_folder(self, session_dir: Path):
        if not self._open_path(session_dir):
            QMessageBox.warning(self, "Dossier", "Impossible d'ouvrir le dossier.")

    def _open_path(self, path: Path) -> bool:
        try:
            p = str(Path(path))
            if os.name == "nt":
                os.startfile(p)
                return True
            if os.name == "posix":
                if "darwin" in (os.uname().sysname or "").lower():
                    subprocess.Popen(["open", p])
                else:
                    subprocess.Popen(["xdg-open", p])
                return True
        except Exception:
            return False
        return False

    def _on_history_item_changed(self, item: QTableWidgetItem):
        if self._history_loading:
            return
        if item.column() != 0:
            return
        session_dir = item.data(Qt.ItemDataRole.UserRole)
        if not session_dir:
            return
        p = Path(str(session_dir)) / "subject.txt"
        try:
            p.write_text(item.text().strip(), encoding="utf-8")
        except Exception:
            pass

    def _start_subject_queue(self):
        if self._subject_worker or not self._subject_queue:
            return
        self._run_next_subject()

    def _run_next_subject(self):
        if self._subject_worker or not self._subject_queue:
            return
        session_dir = self._subject_queue.pop(0)
        transcript_path = self._find_transcript_path(session_dir)
        if not transcript_path:
            return
        self._subject_worker = SubjectWorker(transcript_path, session_dir, self.cfg, parent=self)
        self._subject_worker.finished_ok.connect(self._on_subject_done)
        self._subject_worker.failed.connect(self._on_subject_failed)
        self._subject_worker.start()

    def _on_subject_done(self, session_dir: str, subject: str):
        self._subject_worker = None
        if subject:
            for r in range(self.table_history.rowCount()):
                item = self.table_history.item(r, 0)
                if item and item.data(Qt.ItemDataRole.UserRole) == session_dir:
                    self._history_loading = True
                    item.setText(subject)
                    self._history_loading = False
                    break
        self._run_next_subject()

    def _on_subject_failed(self, session_dir: str, err: str):
        self._subject_worker = None
        self._run_next_subject()

    def _poll_debug_log(self):
        if not self.debug_panel.isVisible():
            return
        try:
            if not LOG_PATH.exists():
                return
            bar = self.debug_panel.verticalScrollBar()
            prev_val = bar.value()
            was_at_bottom = prev_val >= (bar.maximum() - 2)
            data = LOG_PATH.read_text(encoding="utf-8", errors="replace")
            # keep last ~8000 chars to avoid heavy UI updates
            if len(data) > 8000:
                data = data[-8000:]
            self.debug_panel.setPlainText(data)
            if was_at_bottom:
                bar.setValue(bar.maximum())
            else:
                bar.setValue(min(prev_val, bar.maximum()))
        except Exception:
            pass

    def _on_copy_summary(self):
        text = self.txt_summary.toPlainText().strip()
        if not text:
            return
        QGuiApplication.clipboard().setText(text)

    def _on_export_docx(self):
        if self._current_docx and Path(self._current_docx).exists():
            self._open_current_docx()
            return
        self._start_summary_worker(open_docx=True)

    def _on_generate_summary(self):
        self._start_summary_worker(open_docx=False)

    def _on_regenerate_summary(self):
        self._start_summary_worker(open_docx=False)

    def _start_summary_worker(self, open_docx: bool):
        if self._summary_worker:
            return
        if not self._current_session:
            QMessageBox.warning(self, "Résumé", "Sélectionne une session dans l'historique.")
            return
        self.cfg["summary_template"] = str(self.cmb_template.currentText())
        save_config(self.cfg)
        transcript_path = self._find_transcript_path(self._current_session)
        if not transcript_path:
            QMessageBox.warning(self, "Résumé", "Aucune transcription disponible pour générer le DOCX.")
            return

        self._summary_open_after = bool(open_docx)
        self.lbl_progress.setVisible(True)
        self.lbl_progress.setText("Résumé/DOCX en cours...")
        self.pp_progress.setRange(0, 0)  # indeterminate
        self.pp_progress.setVisible(True)

        self._summary_worker = SummaryWorker(transcript_path, self._current_session, self.cfg, parent=self)
        self._summary_worker.finished_ok.connect(self._on_summary_done)
        self._summary_worker.failed.connect(self._on_summary_failed)
        self._summary_worker.start()

    def _on_summary_done(self, docx_path: str, summary_text: str):
        self.pp_progress.setVisible(False)
        self.pp_progress.setRange(0, 100)
        self.lbl_progress.setVisible(False)
        self._summary_worker = None

        if docx_path:
            self._current_docx = Path(docx_path)
            self.btn_open_docx.setEnabled(True)
        if summary_text:
            self.txt_summary.setPlainText(summary_text)
        elif docx_path:
            self.txt_summary.setPlainText("Résumé disponible dans le fichier DOCX (ouvrir pour voir).")
        else:
            self.txt_summary.setPlainText("Résumé non disponible.")

        if self._summary_open_after and docx_path:
            self._open_current_docx()
        self._summary_open_after = False
        if self._current_session:
            self._load_history(focus_session=self._current_session)

    def _on_summary_failed(self, err: str):
        self.pp_progress.setVisible(False)
        self.pp_progress.setRange(0, 100)
        self.lbl_progress.setVisible(False)
        self._summary_worker = None
        self._summary_open_after = False
        QMessageBox.warning(self, "Résumé", f"Erreur génération DOCX: {err}")

    def _on_template_changed(self, text: str):
        self.cfg["summary_template"] = str(text)
        save_config(self.cfg)

    def _apply_diarization_update(self, payload: dict):
        utterances = payload.get("utterances") or []
        if not utterances:
            return
        updated = False
        for item in utterances:
            try:
                s = float(item.get("start", 0.0))
                e = float(item.get("end", 0.0))
                speaker = str(item.get("speaker") or "").strip()
            except Exception:
                continue
            if not speaker:
                continue
            for msg in list(getattr(self.live_tab, "_messages", []) or []):
                if msg.ts_sec >= s and msg.ts_sec <= e:
                    if msg.source != speaker or msg.source_raw != speaker:
                        msg.source_raw = speaker
                        msg.source = self._speaker_display_name(speaker)
                        updated = True

        if updated:
            should_scroll = bool(self.live_tab.auto_scroll and not self.live_tab.scroll_locked)
            try:
                self.live_tab.chat.request_render(keep_scroll=not should_scroll, scroll_to_bottom=should_scroll)
            except Exception:
                pass

    def _apply_cfg_to_ui(self):
        live_enabled = bool(self.cfg.get("enable_live", True))
        common_mod.DEBUG_ENABLED = bool(self.cfg.get("debug_enabled", False))
        self.live_tab.chk_live_enabled.setChecked(live_enabled)
        self.live_tab.btn_rename_speaker.setEnabled(self._is_live_voice_ident_enabled())
        timer_obj = getattr(self, "timer", None)
        is_recording = bool(getattr(self.recorder, "_running", False)) or bool(timer_obj and timer_obj.isActive())
        if is_recording:
            # Preserve runtime controls while a session is running.
            self.topbar.btn_start.setEnabled(False)
            self.topbar.btn_stop.setEnabled(True)
            self.topbar.btn_pause.setEnabled(True)
            self.topbar.pill_audio.set_state("ok")
            if live_enabled:
                self.topbar.pill_transcript.set_state("ok" if self.live_thread else "warn")
            else:
                self.topbar.pill_transcript.set_state("off")
            self.topbar.set_recording(True)
        else:
            self.topbar.btn_start.setEnabled(True)
            self.topbar.btn_stop.setEnabled(False)
            self.topbar.btn_pause.setEnabled(False)
            self.topbar.pill_audio.set_state("off")
            self.topbar.pill_transcript.set_state("off")
            self.topbar.pill_translate.set_state("off")
            self.topbar.pill_summary.set_state("off")
            if live_enabled:
                self.topbar.pill_transcript.set_state("warn")
            self.topbar.set_recording(False)
        part_lang = str(self.cfg.get("live_participant_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        my_lang = str(self.cfg.get("live_my_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        if part_lang not in ("AUTO", "EN", "FR"):
            part_lang = "AUTO"
        if my_lang not in ("AUTO", "EN", "FR"):
            my_lang = "AUTO"
        self._update_live_lang_buttons(part_lang=part_lang, my_lang=my_lang)
        source_role = str(self.cfg.get("live_source_role") or "Participants")
        if source_role.lower() not in ("participants", "moi"):
            source_role = "Participants"
        self._sync_live_language_from_source()

        show_debug = bool(self.cfg.get("debug_show_panel", False)) and bool(self.cfg.get("debug_enabled", False))
        self.debug_panel.setVisible(show_debug)
        if show_debug:
            self._debug_timer.start()
            self._poll_debug_log()
        else:
            self._debug_timer.stop()

    def _save_ui_to_cfg(self):
        self.cfg["enable_live"] = True
        save_config(self.cfg)

    def _update_window_title(self):
        name = self.cfg.get("session_name", "Session")
        dt = datetime.now().strftime("%d-%m")
        self.setWindowTitle(f"{name} - {dt}")

    def _tick_timer(self):
        self._seconds += 1
        hh = self._seconds // 3600
        mm = (self._seconds % 3600) // 60
        ss = self._seconds % 60
        self.topbar.lbl_timer.setText(f"{hh:02d}:{mm:02d}:{ss:02d}")

    def _update_price_estimate(self):
        minutes = max(0.0, float(self._seconds) / 60.0)
        openai_cost = minutes * OPENAI_TRANSCRIBE_PER_MIN
        deepgram_cost = minutes * DEEPGRAM_STREAMING_NOVA3_PER_MIN
        aai_cost = minutes * ASSEMBLYAI_STREAMING_PER_MIN

        def _fmt_usd(v: float) -> str:
            # Display in cents precision (ex: 0,00) as requested.
            return f"{max(0.0, float(v)):.2f}".replace(".", ",")

        engine = str(self.cfg.get("live_engine") or "deepgram").lower()
        if engine == "assemblyai":
            self.lbl_price.setText(
                f"Prix transcription: OpenAI {_fmt_usd(openai_cost)} $ | AssemblyAI {_fmt_usd(aai_cost)} $"
            )
            self.lbl_price.setToolTip(
                "Tarifs indicatifs (USD): OpenAI transcription 0,006 $/min ; AssemblyAI streaming 0,15 $/heure."
            )
        else:
            self.lbl_price.setText(
                f"Prix transcription: OpenAI {_fmt_usd(openai_cost)} $ | Deepgram {_fmt_usd(deepgram_cost)} $"
            )
            self.lbl_price.setToolTip(
                "Tarifs indicatifs (USD): OpenAI transcription 0,006 $/min ; Deepgram Nova-3 streaming 0,0077 $/min."
            )

    def _set_status(self, msg: str):
        self.statusBar().showMessage(msg, 5000)

    def _start_live(self):
        self._cleanup_retained_threads()
        if self.live_thread:
            return

        engine = (self.cfg.get("live_engine") or "deepgram").lower()

        self._sync_live_language_from_source()

        if engine == "deepgram":
            from threads.live_deepgram_thread import LiveDeepgramThread
            self.live_thread = LiveDeepgramThread(cfg=self.cfg, recorder=self.recorder, parent=self)
        else:
            from threads.live_assemblyai_thread import LiveAssemblyAIThread
            self.live_thread = LiveAssemblyAIThread(cfg=self.cfg, recorder=self.recorder, parent=self)

        self.live_thread.live_line.connect(self._append_live)
        self.live_thread.status.connect(self._set_status)
        self.live_thread.start()


    def _stop_live(self):
        t = self.live_thread
        self.live_thread = None
        if not t:
            return
        try:
            t.stop()
        except Exception:
            pass
        try:
            t.quit()
            if not t.wait(4000):
                # Keep a strong ref until the thread actually exits.
                self._retain_thread(t)
        except Exception:
            pass

    def _retain_thread(self, thread: QThread):
        if not thread:
            return
        if thread not in self._retained_threads:
            self._retained_threads.append(thread)
        try:
            thread.finished.connect(lambda _=None, t=thread: self._release_retained_thread(t))
        except Exception:
            pass

    def _release_retained_thread(self, thread: QThread):
        try:
            if thread in self._retained_threads:
                self._retained_threads.remove(thread)
        except Exception:
            pass
        try:
            thread.deleteLater()
        except Exception:
            pass

    def _cleanup_retained_threads(self):
        alive = []
        for t in list(self._retained_threads):
            try:
                if t.isRunning():
                    alive.append(t)
                else:
                    t.deleteLater()
            except Exception:
                pass
        self._retained_threads = alive

    def _on_start(self):
        if self._recording_action_lock:
            return
        if bool(getattr(self.recorder, "_running", False)):
            return
        self._recording_action_lock = True
        try:
            self._chunk_enabled = bool(self.cfg.get("postprocess_chunked", True)) and bool(self.cfg.get("enable_postprocess", False))
            self._reset_chunk_state()
            if self._chunk_enabled:
                self._chunk_cfg_run = dict(self.cfg)
                self._chunk_cfg_start = dict(self.cfg)
                self._chunk_expect_mic = True
                try:
                    self.recorder.clear_part_closed_callbacks()
                    self.recorder.add_part_closed_callback(self._on_part_closed_event)
                except Exception:
                    pass
                self._chunk_poll_timer.start()
            else:
                try:
                    self.recorder.clear_part_closed_callbacks()
                except Exception:
                    pass
            self._configure_live_source_queue()
            self.recorder.start()
            self.session_dir = self.recorder.session_dir

            self._seconds = 0
            self.timer.start(1000)
            self.price_timer.start()
            self._update_price_estimate()
            self._live_start_time = time.time()
            self._live_messages = []
            self._update_window_title()
            self._paused = False
            self.topbar.btn_pause.setText("Pause")

            self.topbar.btn_start.setEnabled(False)
            self.topbar.btn_pause.setEnabled(True)
            self.topbar.btn_stop.setEnabled(True)
            self.topbar.pill_audio.set_state("ok")
            self.topbar.pill_transcript.set_state("ok")
            if bool(self.cfg.get("live_enable_translation", False)):
                self.topbar.pill_translate.set_state("ok")
            else:
                self.topbar.pill_translate.set_state("off")
            if bool(self.cfg.get("enable_postprocess", False)):
                self.topbar.pill_summary.set_state("warn")
            else:
                self.topbar.pill_summary.set_state("off")
            self.topbar.set_recording(True)

            self._set_status("Enregistrement démarré")
            log_line("[UI] Start recording")

            enable_live = bool(self.cfg.get("enable_live", True))
            if enable_live:
                self._start_live()

        except Exception:
            err = traceback.format_exc()
            log_line("=== Start exception ===\n" + err)
            QMessageBox.critical(self, "Erreur", err)
        finally:
            self._recording_action_lock = False

    def _on_pause(self):
        if not self.topbar.btn_pause.isEnabled():
            return
        if not self._paused:
            self._paused = True
            self.timer.stop()
            self.price_timer.stop()
            self.topbar.btn_pause.setText("Reprendre")
            self.topbar.pill_transcript.set_state("warn")
            self._set_status("Pause")
        else:
            self._paused = False
            self.timer.start(1000)
            self.price_timer.start()
            self.topbar.btn_pause.setText("Pause")
            self.topbar.pill_transcript.set_state("ok")
            self._set_status("Reprise")

    def _on_stop(self):
        if self._recording_action_lock:
            return
        timer_obj = getattr(self, "timer", None)
        if not bool(getattr(self.recorder, "_running", False)) and not bool(timer_obj and timer_obj.isActive()):
            return
        self._recording_action_lock = True
        try:
            log_line("[STOP] 1 - click stop")
            self.topbar.btn_stop.setEnabled(False)
            self.timer.stop()
            self.price_timer.stop()
            self._paused = False
            self.topbar.btn_pause.setText("Pause")
            self._live_start_time = None

            log_line("[STOP] 2 - stop live thread")
            self._stop_live()

            log_line("[STOP] 3 - recorder.stop() start")
            self.recorder.stop()
            log_line("[STOP] 4 - recorder.stop() done")
            if self._chunk_poll_timer.isActive():
                self._poll_chunk_events()
                self._chunk_poll_timer.stop()

            self.topbar.btn_start.setEnabled(True)
            self.topbar.btn_pause.setEnabled(False)
            self.topbar.pill_audio.set_state("off")
            self.topbar.pill_transcript.set_state("off")
            self.topbar.pill_translate.set_state("off")
            self.topbar.set_recording(False)

            if not bool(self.cfg.get("enable_postprocess", False)):
                self._set_status("Enregistrement arrêté (compte rendu désactivé)")
                self.topbar.pill_summary.set_state("off")
                return

            if not self.session_dir:
                self.session_dir = self.recorder.session_dir
            if not self.session_dir:
                raise RuntimeError("session_dir manquant (RecorderService)")

            ppaths = list(getattr(self.recorder.participants_track, "wav_paths", []) or [])
            mpaths = list(getattr(self.recorder.my_track, "wav_paths", []) or [])
            self._chunk_expect_mic = bool(mpaths)

            if not ppaths:
                raise RuntimeError("Aucun fichier WAV participants généré.")
            if not mpaths:
                raise RuntimeError("Aucun fichier WAV micro généré.")

            self.recorder.wav_path = Path(ppaths[0])
            self.recorder.mic_wav_path = Path(mpaths[0])

            dlg = StopOptionsDialog(self.cfg, parent=self)
            dlg.exec()
            if dlg.action == "delete":
                self._cancel_chunk_processing(silent=True)
                try:
                    shutil.rmtree(self.session_dir, ignore_errors=True)
                except Exception:
                    pass
                self._reset_ui_state(clear_history=True)
                self._load_history()
                self._set_status("Enregistrement supprimé")
                return
            if dlg.action != "start":
                self._cancel_chunk_processing(silent=True)
                self._set_status("Compte rendu annulé")
                return

            opts = dlg.get_options()
            if not opts.get("generate_transcript", True):
                self._cancel_chunk_processing(silent=True)
                self._set_status("Enregistrement terminé (sans transcription)")
                return

            cfg_run = dict(self.cfg)
            cfg_run.update(
                {
                    "enable_postprocess": True,
                    "postprocess_language": opts.get("postprocess_language", "auto"),
                    "postprocess_quality": opts.get("postprocess_quality", "standard"),
                    "postprocess_device": opts.get("postprocess_device", "auto"),
                    "postprocess_generate_docx": bool(opts.get("generate_report", True)),
                    "postprocess_enable_perplexity_summary": bool(opts.get("postprocess_enable_perplexity_summary", False)),
                    "postprocess_enable_diarization": bool(opts.get("postprocess_enable_diarization", True)),
                    "postprocess_extract_participants": bool(opts.get("postprocess_extract_participants", False)),
                    "summary_template": str(opts.get("summary_template") or "Compte rendu pro"),
                }
            )

            self._set_status("Compte rendu en cours...")
            self.topbar.pill_summary.set_state("warn")

            use_chunk = bool(cfg_run.get("postprocess_chunked", True)) and (
                len(ppaths) > 1 or self._chunk_started or bool(self._chunk_results)
            )
            if use_chunk:
                if self._chunk_started and self._chunk_cfg_run:
                    # Preserve transcription-related settings if chunks already started during recording.
                    for k in (
                        "postprocess_language",
                        "postprocess_quality",
                        "postprocess_device",
                        "postprocess_enable_diarization",
                        "postprocess_diarization_mode",
                    ):
                        if k in self._chunk_cfg_run:
                            cfg_run[k] = self._chunk_cfg_run.get(k)
                else:
                    self._chunk_cfg_run = dict(cfg_run)
                    self._chunk_cfg_start = dict(cfg_run)
                self._chunk_cfg_docx = dict(cfg_run)
                self._start_chunk_postprocess(cfg_run, ppaths, mpaths)
                return

            from threads.postprocess_thread import PostProcessThread

            self.pp_thread = PostProcessThread(cfg=cfg_run, recorder=self.recorder, session_dir=self.session_dir, parent=self)
            self.pp_thread.finished_ok.connect(self._on_postprocess_ok)
            self.pp_thread.failed.connect(self._on_postprocess_fail)
            self.pp_thread.start()
            self._start_postprocess_progress()

        except Exception:
            err = traceback.format_exc()
            log_line("=== Stop exception ===\n" + err)
            QMessageBox.critical(self, "Erreur", err)
            self.topbar.btn_start.setEnabled(True)
        finally:
            self._recording_action_lock = False

    def _on_test_audio(self):
        from ui.setup_window import SetupWindow

        dlg = SetupWindow(start_page="AUDIO")
        dlg.exec()
        self.cfg = load_config()
        try:
            self._reload_recorder_from_cfg()
        except Exception as e:
            QMessageBox.warning(self, "Configuration audio", str(e))
        self._apply_cfg_to_ui()

    def _on_settings(self):
        from ui.setup_window import SetupWindow

        dlg = SetupWindow(start_page="TRANSCRIPTION")
        dlg.exec()
        self.cfg = load_config()
        try:
            self._reload_recorder_from_cfg()
        except Exception as e:
            QMessageBox.warning(self, "Configuration audio", str(e))
        self._apply_cfg_to_ui()

    def _on_live_enabled_changed(self):
        enabled = bool(self.live_tab.chk_live_enabled.isChecked())
        self.cfg["enable_live"] = enabled
        save_config(self.cfg)
        if not enabled:
            self.topbar.pill_transcript.set_state("off")
            self._stop_live()
            self._set_status("Live désactivé")
        else:
            if self.topbar.btn_stop.isEnabled():
                self._start_live()
                self.topbar.pill_transcript.set_state("ok")
            else:
                self.topbar.pill_transcript.set_state("warn")
            self._set_status("Live activé")

    def _on_open_live_chat_window(self):
        if self.live_chat_window is None:
            self.live_chat_window = LiveChatWindow(parent=self)
            self.live_chat_window.set_messages(self._live_messages)
            self.live_chat_window.destroyed.connect(lambda _=None: setattr(self, "live_chat_window", None))
        self.live_chat_window.show()
        self.live_chat_window.raise_()
        self.live_chat_window.activateWindow()

    def _on_reset(self):
        if self.topbar.btn_stop.isEnabled():
            QMessageBox.warning(self, "Reset", "Arrête l'enregistrement avant de reset.")
            return
        self._reset_ui_state(clear_history=False)
        self._set_status("Reset terminé")

    def _reset_ui_state(self, clear_history: bool = False):
        self._last_live_message = None
        self._last_live_turn_id = None
        self._live_messages = []
        self._seconds = 0
        self.topbar.lbl_timer.setText("00:00:00")
        self.live_tab.clear_messages()
        if self.live_chat_window:
            self.live_chat_window.clear_messages()
        if clear_history:
            self.list_transcript.clear()
            self.txt_detail.clear()
            self.txt_summary.clear()
            self.cmb_filter_speaker.clear()
            self.cmb_filter_speaker.addItem("Tous")
        self.topbar.pill_audio.set_state("off")
        self.topbar.pill_transcript.set_state("off")
        self.topbar.pill_translate.set_state("off")
        self.topbar.pill_summary.set_state("off")
        self.topbar.btn_start.setEnabled(True)
        self.topbar.btn_pause.setEnabled(False)
        self.topbar.btn_stop.setEnabled(False)
        self.topbar.set_recording(False)
        self._stop_postprocess_progress()
        self.lbl_cost.setText("Coût: 0 tokens")
        self._update_price_estimate()
        if clear_history:
            self.table_history.setRowCount(0)
            self.txt_summary.setPlainText("")
            self.session_dir = None
        self._reset_chunk_state()

    def _on_cancel_postprocess(self):
        self._cancel_chunk_processing(silent=True)
        if self.pp_thread:
            t = self.pp_thread
            self.pp_thread = None
            try:
                t.stop()
            except Exception:
                pass
            try:
                t.quit()
                if not t.wait(3000):
                    self._retain_thread(t)
            except Exception:
                pass
        self._stop_postprocess_progress()
        self.topbar.pill_summary.set_state("warn")
        self._set_status("Compte rendu annulé")

    def _on_lang_participant_changed(self, text: str):
        self.cfg["live_participant_language"] = str(text).upper()
        if self._sync_live_language_from_source():
            save_config(self.cfg)
            if self.live_thread:
                self._stop_live()
                self._start_live()
        else:
            save_config(self.cfg)
        self._set_status(f"Langue participant: {text}")

    def _on_lang_me_changed(self, text: str):
        self.cfg["live_my_language"] = str(text).upper()
        if self._sync_live_language_from_source():
            save_config(self.cfg)
            if self.live_thread:
                self._stop_live()
                self._start_live()
        else:
            save_config(self.cfg)
        self._set_status(f"Ma langue source: {text}")

    def _on_source_changed(self, text: str):
        self.cfg["live_source_role"] = str(text)
        save_config(self.cfg)
        self._configure_live_source_queue()
        if self.live_thread:
            self._stop_live()
            self._start_live()
        self._set_status(f"Source live appliquée: {text}")

    def _init_live_menus(self):
        menu_part = QMenu(self)
        for label, code in (("Francais", "FR"), ("Anglais", "EN"), ("Automatique", "AUTO")):
            act = QAction(label, self)
            act.triggered.connect(lambda _checked=False, c=code: self._set_participant_lang(c))
            menu_part.addAction(act)
        self.live_tab.btn_lang_participants.setMenu(menu_part)

        menu_me = QMenu(self)
        for label, code in (("Francais", "FR"), ("Anglais", "EN"), ("Automatique", "AUTO")):
            act = QAction(label, self)
            act.triggered.connect(lambda _checked=False, c=code: self._set_my_lang(c))
            menu_me.addAction(act)
        self.live_tab.btn_lang_me.setMenu(menu_me)

    def _set_participant_lang(self, code: str):
        self._update_live_lang_buttons(part_lang=code)
        self._on_lang_participant_changed(code)

    def _set_my_lang(self, code: str):
        self._update_live_lang_buttons(my_lang=code)
        self._on_lang_me_changed(code)

    def _update_live_lang_buttons(self, part_lang: Optional[str] = None, my_lang: Optional[str] = None):
        if part_lang is None:
            part_lang = str(self.cfg.get("live_participant_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        if my_lang is None:
            my_lang = str(self.cfg.get("live_my_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        self.live_tab.btn_lang_participants.setText(self._lang_label(part_lang))
        self.live_tab.btn_lang_participants.setToolTip("Langue du participant")
        self.live_tab.btn_lang_me.setText(self._lang_label(my_lang))
        self.live_tab.btn_lang_me.setToolTip("Ma langue source")

    def _lang_label(self, code: str) -> str:
        code = (code or "").upper()
        if code == "FR":
            return "Français (FR)"
        if code == "EN":
            return "Anglais (US)"
        return "Auto"

    def _sync_live_language_from_source(self) -> bool:
        source_role = str(self.cfg.get("live_source_role") or "Participants").lower()
        if source_role == "moi":
            lang = str(self.cfg.get("live_my_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        else:
            lang = str(self.cfg.get("live_participant_language") or self.cfg.get("live_source_language") or "AUTO").upper()
        if lang not in ("AUTO", "EN", "FR"):
            lang = "AUTO"
        changed = (self.cfg.get("live_source_language") or "").upper() != lang
        self.cfg["live_source_language"] = lang
        return changed

    def _configure_live_source_queue(self):
        source_role = str(self.cfg.get("live_source_role") or "Participants")
        if source_role.lower() not in ("participants", "moi"):
            source_role = "Participants"
        self.live_queue = queue.Queue(maxsize=300)
        self.recorder.live_queue = self.live_queue
        if source_role.lower() == "moi":
            self.recorder.set_live_participants_queue(None)
            self.recorder.set_live_mic_queue(self.live_queue)
        else:
            self.recorder.set_live_mic_queue(None)
            self.recorder.set_live_participants_queue(self.live_queue)
    def _start_postprocess_progress(self):
        self.pp_progress.setValue(0)
        self.pp_progress.setVisible(True)
        self.pp_progress.setRange(0, 100)
        self.lbl_progress.setVisible(True)
        self.lbl_progress.setText("Compte rendu en cours...")
        self.btn_cancel_pp.setVisible(True)
        self._pp_last_payload = None
        self._pp_last_payload_ts = time.time()
        self._pp_display_percent = 0
        self.pp_progress_timer.start()

    def _stop_postprocess_progress(self):
        self.pp_progress_timer.stop()
        self.pp_progress.setVisible(False)
        self.lbl_progress.setVisible(False)
        self.btn_cancel_pp.setVisible(False)
        self._pp_last_payload = None
        self._pp_last_payload_ts = 0.0
        self._pp_display_percent = 0

    def _poll_postprocess_progress(self):
        if not self.session_dir:
            return
        p = Path(self.session_dir) / "postprocess_progress.json"
        if not p.exists():
            return
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            percent = max(0, min(100, int(data.get("percent", 0))))
            msg = data.get("message") or ""
            stage = data.get("stage") or ""
            now = time.time()
            payload = (percent, stage, msg)
            if payload != self._pp_last_payload:
                self._pp_last_payload = payload
                self._pp_last_payload_ts = now
                self._pp_display_percent = max(self._pp_display_percent, percent)
            else:
                # Fallback visual smoothing when backend progress stays static for a while.
                age = max(0.0, now - float(self._pp_last_payload_ts or now))
                if age >= 3.0 and percent < 99:
                    if stage == "start":
                        cap = 25
                    elif stage == "transcription":
                        cap = 88
                    elif stage == "fusion":
                        cap = 94
                    elif stage == "docx":
                        cap = 98
                    else:
                        cap = 90
                    step = int((age - 3.0) // 2.0) + 1
                    synthetic = min(cap, percent + step)
                    self._pp_display_percent = max(self._pp_display_percent, synthetic)

            shown_percent = max(percent, self._pp_display_percent)
            self.pp_progress.setValue(max(0, min(100, shown_percent)))
            if msg or stage:
                stage_map = {
                    "start": "",
                    "prepare": "Préparation",
                    "diarization": "Voix",
                    "transcription": "Transcription",
                    "write": "Finalisation",
                    "fusion": "Fusion",
                    "docx": "DOCX",
                    "done": "Terminé",
                    "error": "Erreur",
                }
                stage_label = stage_map.get(str(stage).strip().lower(), "")
                stage_txt = f"{stage_label} · " if stage_label else ""
                if shown_percent > percent and percent < 100:
                    self.lbl_progress.setText(f"{stage_txt}{msg} · En cours... ({shown_percent}%)")
                else:
                    self.lbl_progress.setText(f"{stage_txt}{msg} ({shown_percent}%)")
        except Exception:
            pass

    def _on_postprocess_ok(self, transcript_path: str):
        self._stop_postprocess_progress()
        self.topbar.btn_start.setEnabled(True)
        self._set_status(f"Terminé: {transcript_path}")
        self.topbar.pill_summary.set_state("ok")
        if self.session_dir:
            self._load_history(focus_session=self.session_dir)
        QMessageBox.information(self, "Terminé", f"Transcription générée:\n{transcript_path}")

    def _on_postprocess_fail(self, err: str):
        if getattr(self, "_suppress_postprocess_errors", False):
            self._suppress_postprocess_errors = False
            self._stop_postprocess_progress()
            self.topbar.btn_start.setEnabled(True)
            self.topbar.pill_summary.set_state("warn")
            self._set_status("Compte rendu annulé")
            return
        self._stop_postprocess_progress()
        self.topbar.btn_start.setEnabled(True)
        self._set_status("Compte rendu: erreur")
        self.topbar.pill_summary.set_state("warn")
        QMessageBox.critical(self, "Compte rendu", err)

    def _append_live(self, line: str):
        text = (line or "").strip()
        if not text:
            return
        if self._paused:
            return

        turn_id: Optional[int] = None
        if text.startswith("[T:"):
            end = text.find("]")
            if end > 3:
                raw = text[3:end]
                try:
                    turn_id = int(raw)
                except Exception:
                    turn_id = None
                text = text[end + 1 :].strip()
        if text.startswith("[DIA]"):
            payload = text.replace("[DIA]", "", 1).strip()
            try:
                data = json.loads(payload)
                self._apply_diarization_update(data)
            except Exception:
                pass
            return
        speaker = None
        if text.startswith("[SPK:"):
            end = text.find("]")
            if end > 5:
                speaker = text[5:end]
                text = text[end + 1 :].strip()
        if text.startswith("[EN]"):
            content = text.replace("[EN]", "", 1).strip()
            self._add_live_message(text_en=content, speaker_override=speaker, turn_id=turn_id)
        elif text.startswith("[FR]"):
            content = text.replace("[FR]", "", 1).strip()
            self._add_live_message(text_fr=content, speaker_override=speaker, turn_id=turn_id)
        else:
            self._add_live_message(text_en=text, speaker_override=speaker, turn_id=turn_id)

    def _is_live_voice_ident_enabled(self) -> bool:
        mode = str(self.cfg.get("voice_identification_mode") or "").strip().lower()
        if mode not in ("report_only", "live_beta"):
            mode = "live_beta" if bool(self.cfg.get("live_speaker_labels", False)) else "report_only"
        return bool(mode == "live_beta" and self.cfg.get("live_speaker_labels", False))

    def _add_live_message(self, text_en: str = "", text_fr: str = "", speaker_override: Optional[str] = None, turn_id: Optional[int] = None):
        def _clean_live_text(t: str) -> str:
            t = (t or "").strip()
            if not t:
                return ""
            # remove zero-width chars
            t = re.sub(r"[\u200b-\u200f\u202a-\u202e]", "", t).strip()
            if not t:
                return ""
            # drop filler-only punctuation
            if re.fullmatch(r"[.\-–—·•…]+", t):
                return ""
            # keep short words if they contain letters/numbers
            if not re.search(r"[\w\d]", t, flags=re.UNICODE):
                return ""
            return t

        def _norm_for_merge(t: str) -> str:
            t = (t or "").strip().lower()
            t = re.sub(r"[^\w\s]", "", t, flags=re.UNICODE)
            t = re.sub(r"\s+", " ", t).strip()
            return t

        text_en = _clean_live_text(text_en)
        text_fr = _clean_live_text(text_fr)
        if not text_en and not text_fr:
            return
        now = datetime.now().strftime("%H:%M")
        if speaker_override and self._is_live_voice_ident_enabled():
            source_raw = speaker_override
        else:
            source_raw = str(self.cfg.get("live_source_role") or "Participants")
            if source_raw.lower() not in ("participants", "moi"):
                source_raw = "Participants"
        source = self._speaker_display_name(source_raw)

        last = self._last_live_message
        last_raw = (last.source_raw if last and last.source_raw else (last.source if last else ""))
        can_merge = bool(last and last_raw == source_raw and (turn_id is None or getattr(self, "_last_live_turn_id", None) == turn_id))
        if can_merge:
            # merge FR after EN
            if text_fr and not last.text_fr and last.text_en:
                last.text_fr = text_fr
                self.live_tab.update_last_message(last)
                self._mirror_live_update(last)
                return
            # update EN partials
            if text_en and last.text_en:
                if text_en == last.text_en:
                    return
                if text_en.startswith(last.text_en) and len(text_en) > len(last.text_en):
                    last.text_en = text_en
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return
                n_new = _norm_for_merge(text_en)
                n_old = _norm_for_merge(last.text_en)
                if n_new.startswith(n_old) and len(n_new) > len(n_old):
                    last.text_en = text_en
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return
                # append short tail fragments
                if turn_id is None and len(text_en.split()) <= 3:
                    last.text_en = (last.text_en.rstrip() + " " + text_en).strip()
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return
            # update FR partials
            if text_fr and last.text_fr:
                if text_fr == last.text_fr:
                    return
                if text_fr.startswith(last.text_fr) and len(text_fr) > len(last.text_fr):
                    last.text_fr = text_fr
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return
                n_new = _norm_for_merge(text_fr)
                n_old = _norm_for_merge(last.text_fr)
                if n_new.startswith(n_old) and len(n_new) > len(n_old):
                    last.text_fr = text_fr
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return
                if turn_id is None and len(text_fr.split()) <= 3:
                    last.text_fr = (last.text_fr.rstrip() + " " + text_fr).strip()
                    self.live_tab.update_last_message(last)
                    self._mirror_live_update(last)
                    return

        # Suppress duplicate text emitted with a different speaker in a very short window.
        if last and last_raw != source_raw:
            try:
                now_sec = 0.0
                if self._live_start_time:
                    now_sec = max(0.0, time.time() - self._live_start_time)
                prev = float(getattr(last, "ts_sec", 0.0) or 0.0)
                if (now_sec - prev) <= 1.2:
                    new_txt = text_fr or text_en
                    last_txt = last.text_fr or last.text_en
                    if _norm_for_merge(new_txt) and _norm_for_merge(new_txt) == _norm_for_merge(last_txt):
                        return
            except Exception:
                pass

        if source and source not in [self.live_tab.cmb_speaker_filter.itemText(i) for i in range(self.live_tab.cmb_speaker_filter.count())]:
            self.live_tab.cmb_speaker_filter.addItem(source)
        ts = 0.0
        if self._live_start_time:
            ts = max(0.0, time.time() - self._live_start_time)
        msg = LiveMessage(timestamp=now, source=source, source_raw=source_raw, ts_sec=ts, text_en=text_en, text_fr=text_fr)
        self._last_live_message = msg
        self._last_live_turn_id = turn_id
        self._live_messages.append(msg)
        self.live_tab.add_message(msg)
        self._mirror_live_append(msg)

    def _speaker_display_name(self, source_raw: str) -> str:
        raw = (source_raw or "").strip()
        if not raw:
            return "Participants"
        alias = (self._live_speaker_aliases.get(raw) or "").strip()
        return alias if alias else raw

    def _on_rename_live_speaker(self):
        if not self._is_live_voice_ident_enabled():
            QMessageBox.information(
                self,
                "Renommer voix",
                "Mode live voix désactivé. Active-le dans Configuration > Transcription > Identification des voix.",
            )
            return
        candidates: list[str] = []
        seen = set()
        for msg in self._live_messages:
            raw = (msg.source_raw or msg.source or "").strip()
            if not raw or raw.lower() in ("participants", "moi"):
                continue
            if raw in seen:
                continue
            seen.add(raw)
            candidates.append(raw)
        if not candidates:
            QMessageBox.information(self, "Renommer voix", "Aucune voix détectée pour le moment.")
            return
        current, ok = QInputDialog.getItem(
            self,
            "Renommer voix",
            "Voix détectée :",
            candidates,
            0,
            False,
        )
        if not ok or not current:
            return
        current_alias = self._live_speaker_aliases.get(current, "")
        new_name, ok = QInputDialog.getText(
            self,
            "Renommer voix",
            f"Nouveau nom pour '{current}' :",
            text=current_alias or current,
        )
        if not ok:
            return
        new_name = (new_name or "").strip()
        if not new_name:
            self._live_speaker_aliases.pop(current, None)
        else:
            self._live_speaker_aliases[current] = new_name
        self.cfg["live_speaker_aliases"] = dict(self._live_speaker_aliases)
        save_config(self.cfg)
        self._refresh_live_speaker_labels()

    def _refresh_live_speaker_labels(self):
        for msg in self._live_messages:
            raw = (msg.source_raw or msg.source or "").strip()
            msg.source = self._speaker_display_name(raw)
        if self._last_live_message:
            raw = (self._last_live_message.source_raw or self._last_live_message.source or "").strip()
            self._last_live_message.source = self._speaker_display_name(raw)
        self.live_tab.clear_messages()
        for msg in self._live_messages:
            self.live_tab.add_message(msg)
        if self.live_chat_window:
            self.live_chat_window.set_messages(self._live_messages)

    def _on_copy_last_live(self):
        msg = self.live_tab.get_last_message()
        if not msg:
            return
        txt = ""
        if msg.text_en:
            txt += f"EN: {msg.text_en}\n"
        if msg.text_fr:
            txt += f"FR: {msg.text_fr}\n"
        QGuiApplication.clipboard().setText(txt.strip())

    def _on_mark_important_live(self):
        msg = self.live_tab.get_last_message()
        if not msg:
            return
        msg.important = not msg.important
        self.live_tab.update_last_message(msg)
        self._mirror_live_update(msg)

    def _mirror_live_append(self, msg: LiveMessage):
        if not self.live_chat_window:
            return
        try:
            self.live_chat_window.add_message(msg, keep_scroll=False)
        except Exception:
            pass

    def _mirror_live_update(self, msg: LiveMessage):
        if not self.live_chat_window:
            return
        try:
            self.live_chat_window.update_last_message(msg, keep_scroll=False)
        except Exception:
            pass

    def _reset_chunk_state(self, init: bool = False):
        self._chunk_part_files = {}
        self._chunk_queue = []
        self._chunk_enqueued = set()
        self._chunk_results = {}
        self._chunk_expected_parts = set()
        self._chunk_started = False
        self._chunk_finalize_requested = False
        self._chunk_cancelled = False
        self._chunk_cfg_run = None
        self._chunk_cfg_start = None
        self._chunk_cfg_docx = None
        self._chunk_thread = None
        self._chunk_merge_thread = None
        self._chunk_expect_mic = True
        self._chunk_event_queue = queue.Queue()
        if not init:
            try:
                self._chunk_poll_timer.stop()
            except Exception:
                pass
            try:
                self.recorder.clear_part_closed_callbacks()
            except Exception:
                pass

    def _on_part_closed_event(self, label: str, path: Path):
        try:
            self._chunk_event_queue.put_nowait((str(label), str(path)))
        except Exception:
            pass

    def _poll_chunk_events(self):
        if not self._chunk_enabled or self._chunk_cancelled:
            return
        while True:
            try:
                label, path = self._chunk_event_queue.get_nowait()
            except queue.Empty:
                break
            try:
                self._register_part_file(label, Path(path))
            except Exception:
                pass
        self._start_next_chunk()

    def _extract_part_index(self, path: Path) -> Optional[int]:
        name = str(Path(path).name)
        m = re.search(r"partie(\d+)", name, flags=re.IGNORECASE)
        if not m:
            return None
        try:
            return int(m.group(1))
        except Exception:
            return None

    def _register_part_file(self, label: str, path: Path):
        if not path:
            return
        idx = self._extract_part_index(path)
        if idx is None:
            return
        entry = self._chunk_part_files.setdefault(idx, {})
        if label == getattr(self.recorder, "participants_label", None):
            entry["participants"] = Path(path)
        elif label == getattr(self.recorder, "my_audio_label", None):
            entry["mic"] = Path(path)
        else:
            lower = path.name.lower()
            if "participants" in lower:
                entry["participants"] = Path(path)
            elif "micro" in lower or "mic" in lower:
                entry["mic"] = Path(path)
        ready = "participants" in entry and (not self._chunk_expect_mic or "mic" in entry)
        if ready:
            self._enqueue_chunk(idx)

    def _enqueue_chunk(self, idx: int):
        if idx in self._chunk_enqueued or idx in self._chunk_results:
            return
        self._chunk_enqueued.add(idx)
        self._chunk_queue.append(idx)

    def _chunk_total(self) -> int:
        if self._chunk_expected_parts:
            return max(1, len(self._chunk_expected_parts))
        pending = len(self._chunk_queue)
        running = 1 if self._chunk_thread else 0
        return max(1, len(self._chunk_results) + pending + running)

    def _update_chunk_progress(self, stage: str, message: str):
        total = self._chunk_total()
        done = len(self._chunk_results)
        percent = 10 + int(60 * (float(done) / float(total)))
        percent = max(5, min(75, percent))
        self._write_progress(percent, stage, message)

    def _start_next_chunk(self):
        if not self._chunk_enabled or self._chunk_cancelled:
            return
        if self._chunk_thread or not self._chunk_queue:
            return
        if not self.session_dir:
            return
        idx = self._chunk_queue.pop(0)
        entry = self._chunk_part_files.get(idx, {})
        wav_path = entry.get("participants")
        if not wav_path:
            return
        mic_path = entry.get("mic")

        chunk_dir = Path(self.session_dir) / "chunks" / f"part_{idx:02d}"
        chunk_dir.mkdir(parents=True, exist_ok=True)

        cfg = self._chunk_cfg_run or self._chunk_cfg_start or dict(self.cfg)
        from threads.postprocess_thread import ChunkPostProcessThread

        self._chunk_thread = ChunkPostProcessThread(
            cfg=cfg,
            part_index=idx,
            wav_path=wav_path,
            mic_wav_path=mic_path,
            session_dir=chunk_dir,
        )
        self._chunk_thread.finished_ok.connect(self._on_chunk_finished)
        self._chunk_thread.failed.connect(self._on_chunk_failed)
        self._chunk_thread.start()
        self._chunk_started = True
        self._update_chunk_progress("transcription", f"Segment {idx}/{self._chunk_total()}")

    def _on_chunk_finished(self, part_index: int, transcript_path: str):
        self._chunk_thread = None
        self._chunk_results[int(part_index)] = str(transcript_path)
        self._update_chunk_progress("transcription", f"Segment {len(self._chunk_results)}/{self._chunk_total()} terminé")
        self._start_next_chunk()
        self._maybe_finalize_chunks()

    def _on_chunk_failed(self, part_index: int, err: str):
        self._chunk_thread = None
        self._chunk_cancelled = True
        self._write_progress(100, "error", f"Erreur segment {part_index}")
        if getattr(self, "_suppress_postprocess_errors", False):
            return
        self._on_postprocess_fail(err)

    def _start_chunk_postprocess(self, cfg_run: dict, ppaths: list[Path], mpaths: list[Path]):
        if not self.session_dir:
            return
        self._chunk_finalize_requested = True
        self._chunk_cfg_docx = dict(cfg_run)
        part_indices = set()
        for p in ppaths:
            idx = self._extract_part_index(Path(p))
            if idx is not None:
                part_indices.add(idx)
        self._chunk_expected_parts = part_indices

        # Ensure all parts are registered (useful after stop when callbacks are done).
        for p in ppaths:
            self._register_part_file(self.recorder.participants_label, Path(p))
        for m in mpaths:
            self._register_part_file(self.recorder.my_audio_label, Path(m))

        self._start_postprocess_progress()
        self._write_progress(5, "start", "Compte rendu démarré")
        self._start_next_chunk()
        self._maybe_finalize_chunks()

    def _maybe_finalize_chunks(self):
        if not self._chunk_finalize_requested or self._chunk_cancelled:
            return
        if self._chunk_thread or self._chunk_queue:
            return
        if self._chunk_expected_parts and len(self._chunk_results) < len(self._chunk_expected_parts):
            return
        self._start_chunk_merge()

    def _start_chunk_merge(self):
        if self._chunk_merge_thread or self._chunk_cancelled:
            return
        if not self.session_dir:
            return
        parts = []
        target_indices = sorted(self._chunk_expected_parts or self._chunk_results.keys())
        for idx in target_indices:
            entry = self._chunk_part_files.get(idx, {})
            transcript_path = self._chunk_results.get(idx)
            wav_path = entry.get("participants")
            if not transcript_path or not wav_path:
                self._chunk_cancelled = True
                self._on_postprocess_fail(f"Segment manquant pour Partie{idx:02d}")
                return
            parts.append(
                {
                    "index": idx,
                    "transcript_path": transcript_path,
                    "participants_wav": str(wav_path),
                }
            )

        self._chunk_merge_thread = ChunkMergeWorker(
            session_dir=self.session_dir,
            parts=parts,
            cfg=self._chunk_cfg_docx or self.cfg,
            generate_docx=bool((self._chunk_cfg_docx or {}).get("postprocess_generate_docx", True)),
        )
        self._chunk_merge_thread.finished_ok.connect(self._on_chunk_merge_ok)
        self._chunk_merge_thread.failed.connect(self._on_chunk_merge_fail)
        self._chunk_merge_thread.start()

    def _on_chunk_merge_ok(self, transcript_path: str, docx_path: str):
        self._chunk_merge_thread = None
        self._on_postprocess_ok(transcript_path)

    def _on_chunk_merge_fail(self, err: str):
        self._chunk_merge_thread = None
        if getattr(self, "_suppress_postprocess_errors", False):
            return
        self._on_postprocess_fail(err)

    def _write_progress(self, percent: int, stage: str, message: str):
        if not self.session_dir:
            return
        try:
            payload = {
                "percent": int(percent),
                "stage": str(stage),
                "message": str(message),
            }
            (Path(self.session_dir) / "postprocess_progress.json").write_text(
                json.dumps(payload, ensure_ascii=False), encoding="utf-8"
            )
        except Exception:
            pass

    def _cancel_chunk_processing(self, silent: bool = False):
        self._chunk_cancelled = True
        if silent:
            self._suppress_postprocess_errors = True
        try:
            self._chunk_poll_timer.stop()
        except Exception:
            pass
        if self._chunk_thread:
            try:
                self._chunk_thread.stop()
            except Exception:
                pass
            try:
                self._chunk_thread.quit()
                self._chunk_thread.wait(2000)
            except Exception:
                pass
            self._chunk_thread = None
        if self._chunk_merge_thread:
            try:
                self._chunk_merge_thread.requestInterruption()
            except Exception:
                pass
            try:
                self._chunk_merge_thread.quit()
                self._chunk_merge_thread.wait(2000)
            except Exception:
                pass
            self._chunk_merge_thread = None
        self._chunk_queue = []
        self._chunk_enqueued = set()
        self._chunk_expected_parts = set()
        self._chunk_results = {}
        self._chunk_finalize_requested = False
        try:
            self.recorder.clear_part_closed_callbacks()
        except Exception:
            pass

    def closeEvent(self, event):
        # Ensure no worker thread survives the window close.
        try:
            self._cancel_chunk_processing(silent=True)
        except Exception:
            pass
        try:
            if self.pp_thread:
                t = self.pp_thread
                self.pp_thread = None
                try:
                    t.stop()
                except Exception:
                    pass
                try:
                    t.quit()
                    t.wait(3000)
                except Exception:
                    pass
        except Exception:
            pass
        try:
            self._stop_live()
        except Exception:
            pass
        try:
            if bool(getattr(self.recorder, "_running", False)):
                self.recorder.stop()
        except Exception:
            pass
        self._cleanup_retained_threads()
        return super().closeEvent(event)


class SummaryWorker(QThread):
    finished_ok = pyqtSignal(str, str)  # docx_path, summary_text
    failed = pyqtSignal(str)

    def __init__(self, transcript_path: Path, session_dir: Path, cfg: dict, parent=None):
        super().__init__(parent)
        self.transcript_path = Path(transcript_path)
        self.session_dir = Path(session_dir)
        self.cfg = cfg or {}

    def run(self) -> None:
        try:
            from services.meeting_summary_service import generate_meeting_docx

            docx_path = generate_meeting_docx(
                transcript_path=self.transcript_path,
                session_dir=self.session_dir,
                cfg=self.cfg,
            )
            summary_text = ""
            try:
                from docx import Document

                doc = Document(str(docx_path))
                summary_text = "\n".join(p.text for p in doc.paragraphs if p.text).strip()
            except Exception:
                summary_text = ""
            self.finished_ok.emit(str(docx_path), summary_text)
        except Exception as e:
            self.failed.emit(repr(e))


class ChunkMergeWorker(QThread):
    finished_ok = pyqtSignal(str, str)  # transcript_path, docx_path
    failed = pyqtSignal(str)

    def __init__(self, session_dir: Path, parts: list[dict], cfg: dict, generate_docx: bool, parent=None):
        super().__init__(parent)
        self.session_dir = Path(session_dir)
        self.parts = list(parts)
        self.cfg = cfg or {}
        self.generate_docx = bool(generate_docx)

    def _write_progress(self, percent: int, stage: str, message: str) -> None:
        try:
            payload = {
                "percent": int(percent),
                "stage": str(stage),
                "message": str(message),
            }
            (self.session_dir / "postprocess_progress.json").write_text(
                json.dumps(payload, ensure_ascii=False), encoding="utf-8"
            )
        except Exception:
            pass

    def _wav_duration_seconds(self, path: Path) -> float:
        try:
            with wave.open(str(path), "rb") as wf:
                frames = wf.getnframes()
                rate = wf.getframerate() or 1
                return float(frames) / float(rate)
        except Exception:
            return 0.0

    def _parse_ts(self, s: str) -> float:
        try:
            parts = s.split(":")
            if len(parts) != 3:
                return 0.0
            h, m, sec = [int(p) for p in parts]
            return float(h * 3600 + m * 60 + sec)
        except Exception:
            return 0.0

    def _fmt_ts(self, t: float) -> str:
        if t < 0:
            t = 0.0
        h = int(t // 3600)
        m = int((t % 3600) // 60)
        s = int(t % 60)
        return f"{h:02d}:{m:02d}:{s:02d}"

    def _merge_lines(self, lines: list[str], offset: float) -> list[str]:
        out = []
        pattern = re.compile(r"^\[(?P<start>\d{2}:\d{2}:\d{2})\s*-\s*(?P<end>\d{2}:\d{2}:\d{2})\]\s*(?P<rest>.*)$")
        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            m = pattern.match(line)
            if not m:
                out.append(line)
                continue
            start_s = self._parse_ts(m.group("start")) + offset
            end_s = self._parse_ts(m.group("end")) + offset
            rest = m.group("rest").strip()
            out.append(f"[{self._fmt_ts(start_s)} - {self._fmt_ts(end_s)}] {rest}".strip())
        return out

    def run(self) -> None:
        try:
            self._write_progress(80, "fusion", "Fusion des segments")
            merged_lines: list[str] = []
            offset = 0.0
            parts_sorted = sorted(self.parts, key=lambda p: int(p.get("index", 0)))
            for part in parts_sorted:
                if self.isInterruptionRequested():
                    raise RuntimeError("Compte rendu annulé")
                tpath = Path(part.get("transcript_path"))
                wav_path = Path(part.get("participants_wav"))
                try:
                    lines = tpath.read_text(encoding="utf-8").splitlines()
                except Exception:
                    lines = []
                merged_lines.extend(self._merge_lines(lines, offset))
                offset += self._wav_duration_seconds(wav_path)

            out_mix = self.session_dir / "transcript-speakers.mix.txt"
            out_fr = self.session_dir / "transcript_speakers_fr.txt"
            content = "\n".join(merged_lines) + ("\n" if merged_lines else "")
            out_mix.write_text(content, encoding="utf-8")
            out_fr.write_text(content, encoding="utf-8")

            docx_path = ""
            if self.generate_docx:
                self._write_progress(92, "docx", "Génération du DOCX...")
                from services.meeting_summary_service import generate_meeting_docx

                out = generate_meeting_docx(out_mix, self.session_dir, self.cfg)
                docx_path = str(out)
                self._write_progress(100, "done", "DOCX généré")
            else:
                self._write_progress(100, "done", "Transcription terminée")

            result = {"transcript_path": str(out_mix), "docx_path": docx_path}
            (self.session_dir / "postprocess_result.json").write_text(
                json.dumps(result, ensure_ascii=False), encoding="utf-8"
            )
            self.finished_ok.emit(str(out_mix), docx_path)
        except Exception:
            self._write_progress(100, "error", "Compte rendu en erreur")
            self.failed.emit(traceback.format_exc())
